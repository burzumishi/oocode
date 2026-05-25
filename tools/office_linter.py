#!/usr/bin/env python3
"""Office format linter — backend de efm-langserver para OOCode.

Uso: office_linter.py <filepath>
Salida: una línea por diagnóstico en formato efm-langserver:
    /ruta/al/fichero:line:col: severity: message

Formatos soportados:
  .docx .doc .dotx .docm  — python-docx
  .xlsx .xlsm .xltx       — openpyxl
  .xls                    — legacy Excel (aviso de conversión)
  .csv                    — csv stdlib
  .pdf                    — pdftotext (poppler-utils)
  .odt .ods .odp          — zipfile + content.xml
"""
import re
import subprocess
import sys
from pathlib import Path


def _diag(path: str, line: int, col: int, sev: str, msg: str) -> str:
    return f"{path}:{line}:{col}: {sev}: {msg}"


def _template_fields(text: str) -> list[str]:
    return list(set(re.findall(r"\{\{([A-Z_a-z][A-Z_a-z0-9]*)\}\}", text)))


# ── DOCX ─────────────────────────────────────────────────────────────────────

def lint_docx(path: Path) -> list[str]:
    diags: list[str] = []
    p = str(path)
    try:
        from docx import Document  # type: ignore
        doc = Document(p)
        for i, para in enumerate(doc.paragraphs, 1):
            for f in _template_fields(para.text):
                diags.append(_diag(p, i, 1, "warning", f"Campo sin rellenar: {{{{{f}}}}}"))
        for tbl_i, tbl in enumerate(doc.tables, 1):
            for row in tbl.rows:
                for cell in row.cells:
                    for f in _template_fields(cell.text):
                        diags.append(_diag(p, tbl_i, 1, "warning",
                                           f"Campo sin rellenar en tabla: {{{{{f}}}}}"))
        raw = path.read_bytes()
        if b"w:ins " in raw or b"w:del " in raw:
            diags.append(_diag(p, 1, 1, "information",
                               "Documento con cambios sin aceptar (Track Changes activo)"))
    except ImportError:
        diags.append(_diag(p, 1, 1, "warning",
                           "python-docx no instalado: pip install python-docx"))
    except Exception as exc:
        diags.append(_diag(p, 1, 1, "error", f"No se puede leer el DOCX: {exc}"))
    return diags


# ── XLSX ─────────────────────────────────────────────────────────────────────

_FORMULA_ERRORS = frozenset({"#REF!", "#VALUE!", "#NAME?", "#DIV/0!", "#N/A", "#NUM!", "#NULL!"})


def lint_xlsx(path: Path) -> list[str]:
    diags: list[str] = []
    p = str(path)
    try:
        import openpyxl  # type: ignore
        wb = openpyxl.load_workbook(p, data_only=False)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    v = str(cell.value)
                    if v in _FORMULA_ERRORS:
                        diags.append(_diag(p, cell.row, cell.column, "error",
                                           f"{sheet_name}!{cell.coordinate}: error de fórmula {v}"))
                    for f in _template_fields(v):
                        diags.append(_diag(p, cell.row, cell.column, "warning",
                                           f"{sheet_name}!{cell.coordinate}: campo sin rellenar {{{{{f}}}}}"))
    except ImportError:
        diags.append(_diag(p, 1, 1, "warning",
                           "openpyxl no instalado: pip install openpyxl"))
    except Exception as exc:
        diags.append(_diag(p, 1, 1, "error", f"No se puede leer el XLSX: {exc}"))
    return diags


# ── XLS (legacy) ─────────────────────────────────────────────────────────────

def lint_xls(path: Path) -> list[str]:
    p = str(path)
    diags = [_diag(p, 1, 1, "information",
                   "Formato .xls legado — considera convertir a .xlsx para soporte completo")]
    try:
        import openpyxl  # type: ignore
        wb = openpyxl.load_workbook(p, data_only=False)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and str(cell.value) in _FORMULA_ERRORS:
                        diags.append(_diag(p, cell.row, cell.column, "error",
                                           f"{sheet_name}!{cell.coordinate}: error de fórmula {cell.value}"))
    except Exception:
        pass
    return diags


# ── CSV ──────────────────────────────────────────────────────────────────────

def lint_csv(path: Path) -> list[str]:
    import csv
    diags: list[str] = []
    p = str(path)
    try:
        with path.open(newline="", errors="replace") as f:
            reader = csv.reader(f)
            try:
                headers = next(reader)
            except StopIteration:
                diags.append(_diag(p, 1, 1, "warning", "CSV vacío"))
                return diags
            if headers and headers[0].startswith("﻿"):
                diags.append(_diag(p, 1, 1, "warning",
                                   "CSV contiene BOM (Byte Order Mark) — puede causar problemas de parseo"))
            num_cols = len(headers)
            for i, row in enumerate(reader, 2):
                if len(row) != num_cols:
                    diags.append(_diag(p, i, 1, "error",
                                       f"Número de columnas inconsistente: esperado {num_cols}, "
                                       f"encontrado {len(row)}"))
    except Exception as exc:
        diags.append(_diag(p, 1, 1, "error", f"Error leyendo CSV: {exc}"))
    return diags


# ── PDF ──────────────────────────────────────────────────────────────────────

def lint_pdf(path: Path) -> list[str]:
    diags: list[str] = []
    p = str(path)
    try:
        r = subprocess.run(
            ["pdftotext", p, "-"],
            capture_output=True, text=True, timeout=15
        )
        if r.returncode != 0:
            msg = (r.stderr.strip() or "error desconocido")[:100]
            diags.append(_diag(p, 1, 1, "error", f"PDF no legible o protegido: {msg}"))
            return diags
        text = r.stdout
        if not text.strip():
            diags.append(_diag(p, 1, 1, "warning",
                               "PDF sin texto extraíble (imagen escaneada o protegido)"))
        for f in _template_fields(text):
            diags.append(_diag(p, 1, 1, "warning",
                               f"Campo sin rellenar en PDF: {{{{{f}}}}}"))
    except FileNotFoundError:
        diags.append(_diag(p, 1, 1, "warning",
                           "pdftotext no disponible: apt install poppler-utils"))
    except subprocess.TimeoutExpired:
        diags.append(_diag(p, 1, 1, "warning", "PDF demasiado grande — análisis cancelado (15s timeout)"))
    except Exception as exc:
        diags.append(_diag(p, 1, 1, "error", f"Error validando PDF: {exc}"))
    return diags


# ── ODT / ODS / ODP ──────────────────────────────────────────────────────────

def lint_odt(path: Path) -> list[str]:
    import zipfile
    diags: list[str] = []
    p = str(path)
    try:
        with zipfile.ZipFile(p) as z:
            content = z.read("content.xml").decode("utf-8", errors="replace")
        text = re.sub(r"<[^>]+>", " ", content)
        for f in _template_fields(text):
            diags.append(_diag(p, 1, 1, "warning", f"Campo sin rellenar: {{{{{f}}}}}"))
    except KeyError:
        diags.append(_diag(p, 1, 1, "error",
                           "Fichero ODF inválido: content.xml no encontrado"))
    except Exception as exc:
        diags.append(_diag(p, 1, 1, "error", f"No se puede leer el fichero ODF: {exc}"))
    return diags


# ── Main ─────────────────────────────────────────────────────────────────────

_LINTERS = {
    ".docx": lint_docx, ".doc": lint_docx, ".dotx": lint_docx, ".docm": lint_docx,
    ".xlsx": lint_xlsx, ".xlsm": lint_xlsx, ".xltx": lint_xlsx,
    ".xls":  lint_xls,
    ".csv":  lint_csv,
    ".pdf":  lint_pdf,
    ".odt":  lint_odt, ".ods":  lint_odt, ".odp": lint_odt, ".odg": lint_odt,
}


def main() -> None:
    if len(sys.argv) < 2:
        sys.exit(0)
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"{path}:1:1: error: Fichero no encontrado: {path.name}")
        sys.exit(0)
    linter = _LINTERS.get(path.suffix.lower())
    if linter is None:
        sys.exit(0)
    for diag in linter(path):
        print(diag)


if __name__ == "__main__":
    main()
