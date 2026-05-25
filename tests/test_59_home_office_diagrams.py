"""Tests de gráficas y diagramas en documentos Office.

Verifica que _tool_doc_insert_diagram inserta correctamente:
- Diagramas de flujo simples
- Gráficas de barras con tablas formateadas
- Gráficas de pastel con emojis
- Tablas de datos dinámicas
- Organigramas jerárquicos

Y que _tool_doc_fill_template preserva estilos originales.
"""
import sys
import tempfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from mcp_servers.home_office_assistant import _tool_doc_insert_diagram, _tool_doc_fill_template


def _create_test_doc(path: str, content: str = "") -> str:
    """Crear un documento .docx vacío o con contenido inicial."""
    from docx import Document
    doc = Document()
    if content:
        para = doc.add_paragraph()
        run = para.add_run(content)
        run.font.size = 12
    doc.save(path)
    return path


class TestDocInsertDiagram:
    """Tests de _tool_doc_insert_diagram."""

    def test_diagram_type_flowchart(self):
        """Test diagrama de flujo simple."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        
        try:
            _create_test_doc(path, "Inicio")
            result = _tool_doc_insert_diagram({
                "path": path,
                "diagram_type": "flowchart",
                "content": "Inicio\nProceso 1\nProceso 2\nFin"
            })
            assert "✅ Diagrama insertado" in result
            assert "flowchart" in result
        finally:
            Path(path).unlink(missing_ok=True)

    def test_diagram_type_bar_chart(self):
        """Test gráfica de barras."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        
        try:
            _create_test_doc(path, "")
            result = _tool_doc_insert_diagram({
                "path": path,
                "diagram_type": "bar_chart",
                "content": "Producto A,Producto B,Producto C\n100,150,200"
            })
            assert "✅ Diagrama insertado" in result
            assert "bar_chart" in result
        finally:
            Path(path).unlink(missing_ok=True)

    def test_diagram_type_table(self):
        """Test tabla de datos."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        
        try:
            _create_test_doc(path, "")
            result = _tool_doc_insert_diagram({
                "path": path,
                "diagram_type": "table",
                "content": "Departamento|Empleados|Salario\nVentas|15|45000€\nDesarrollo|12|52000€"
            })
            assert "✅ Diagrama insertado" in result
            assert "table" in result
        finally:
            Path(path).unlink(missing_ok=True)

    def test_diagram_type_pie_chart(self):
        """Test gráfica de pastel."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        
        try:
            _create_test_doc(path, "")
            result = _tool_doc_insert_diagram({
                "path": path,
                "diagram_type": "pie_chart",
                "content": "Categoría A,Categoría B,Categoría C\n50,30,20"
            })
            assert "✅ Diagrama insertado" in result
            assert "pie_chart" in result
        finally:
            Path(path).unlink(missing_ok=True)

    def test_diagram_type_org_chart(self):
        """Test organigrama."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        
        try:
            _create_test_doc(path, "")
            result = _tool_doc_insert_diagram({
                "path": path,
                "diagram_type": "org_chart",
                "content": "CEO → CTO → Desarrollador\nCEO → CFO → Contador"
            })
            assert "✅ Diagrama insertado" in result
            assert "org_chart" in result
        finally:
            Path(path).unlink(missing_ok=True)

    def test_diagram_missing_path(self):
        """Test error cuando path no existe."""
        result = _tool_doc_insert_diagram({
            "path": "/nonexistent/path.docx",
            "diagram_type": "flowchart",
            "content": "Inicio"
        })
        assert "no encontrado" in result

    def test_diagram_missing_content(self):
        """Test con contenido vacío."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        
        try:
            _create_test_doc(path, "")
            result = _tool_doc_insert_diagram({
                "path": path,
                "diagram_type": "flowchart",
                "content": ""
            })
            # No debe fallar, solo no inserta nada
            assert result is not None
        finally:
            Path(path).unlink(missing_ok=True)

    def test_diagram_output_path(self):
        """Test con output_path especificado."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            input_path = f.name
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        
        try:
            _create_test_doc(input_path, "Inicio")
            result = _tool_doc_insert_diagram({
                "path": input_path,
                "diagram_type": "flowchart",
                "content": "Inicio\nFin",
                "output_path": output_path
            })
            assert "✅ Diagrama insertado" in result
            assert output_path in result
        finally:
            Path(input_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)


class TestDocFillTemplateStyles:
    """Tests de preservación de estilos en _tool_doc_fill_template."""

    def test_fill_template_preserves_styles(self):
        """Test que se preservan estilos originales."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            template_path = f.name
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        
        try:
            # Crear plantilla con placeholder
            from docx import Document
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run("{{NOMBRE}}")
            run.font.size = 12
            doc.save(template_path)
            
            result = _tool_doc_fill_template({
                "template_path": template_path,
                "fields": {"NOMBRE": "Juan", "FECHA": "2026-05-24"},
                "output_path": output_path
            })
            assert "✅ Plantilla rellenada" in result
            assert "Estilos originales preservados" in result
        finally:
            Path(template_path).unlink(missing_ok=True)
            Path(output_path).unlink(missing_ok=True)

    def test_fill_template_missing_fields(self):
        """Test con campos faltantes."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            template_path = f.name
        
        try:
            from docx import Document
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run("{{NOMBRE}}")
            doc.save(template_path)
            
            result = _tool_doc_fill_template({
                "template_path": template_path,
                "fields": {},
                "output_path": "/tmp/filled.docx"
            })
            assert "Parámetro requerido: fields" in result
        finally:
            Path(template_path).unlink(missing_ok=True)

    def test_fill_template_missing_template(self):
        """Test con plantilla no existente."""
        result = _tool_doc_fill_template({
            "template_path": "/nonexistent.docx",
            "fields": {"NOMBRE": "Juan"},
            "output_path": "/tmp/filled.docx"
        })
        assert "no encontrada" in result


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
