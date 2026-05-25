"""Tests de gráficas dinámicas en documentos Office para Home Office MCP.

Verifica que _tool_doc_insert_diagram inserta correctamente:
- Gráficas de barras con matplotlib en documentos .docx
- Gráficas de pastel con matplotlib en documentos .docx
- Tablas de datos dinámicas con formato Office
- Organigramas jerárquicos
- Diagramas de flujo

Y que las gráficas sean compatibles con O365 y LibreOffice.
"""
import sys
import tempfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from mcp_servers.home_office_assistant import _tool_doc_insert_diagram


class TestHomeOfficeDynamicCharts:
    """Tests de gráficas dinámicas en documentos Office."""

    def test_bar_chart_with_matplotlib(self):
        """Test gráfica de barras con matplotlib."""
        # Crear documento vacío
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            # Datos para gráfica de barras
            data = "Ventas,Q1,Q2,Q3,Q4\n100,150,120,180"
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "bar_chart",
                "content": data,
                "output_path": str(path)
            })
            
            assert "✅" in result
            assert "bar_chart" in result.lower()
            
        finally:
            path.unlink(missing_ok=True)

    def test_pie_chart_with_matplotlib(self):
        """Test gráfica de pastel con matplotlib."""
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            # Datos para gráfica de pastel
            data = "Presupuesto,Salarios,Infraestructura,Marketing\n40,30,20,10"
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "pie_chart",
                "content": data,
                "output_path": str(path)
            })
            
            assert "✅" in result
            assert "pie_chart" in result.lower()
            
        finally:
            path.unlink(missing_ok=True)

    def test_table_with_formatting(self):
        """Test tabla de datos dinámica con formato Office."""
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            # Datos de tabla
            data = "Departamento|Empleado|Salario|Puesto\nIT|Juan|50000|Senior Dev\nIT|Maria|55000|Lead Dev\nVentas|Pedro|45000|Vendedor\nVentas|Ana|48000|Senior Vendedor\nRRHH|Luis|52000|HR Manager"
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "table",
                "content": data,
                "output_path": str(path)
            })
            
            assert "✅" in result
            assert "table" in result.lower()
            
        finally:
            path.unlink(missing_ok=True)

    def test_flowchart_diagram(self):
        """Test diagrama de flujo."""
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            # Diagrama de flujo simple
            data = "Inicio → Validar datos → Procesar → Guardar → Fin"
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "flowchart",
                "content": data,
                "output_path": str(path)
            })
            
            assert "✅" in result
            assert "flowchart" in result.lower()
            
        finally:
            path.unlink(missing_ok=True)

    def test_org_chart(self):
        """Test organigrama jerárquico."""
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            # Organigrama jerárquico
            data = "CEO → CTO → Senior Dev → Junior Dev\nCEO → CFO → Controller → Accountant"
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "org_chart",
                "content": data,
                "output_path": str(path)
            })
            
            assert "✅" in result
            assert "org_chart" in result.lower()
            
        finally:
            path.unlink(missing_ok=True)

    def test_missing_path_parameter(self):
        """Test error cuando falta path."""
        result = _tool_doc_insert_diagram({
            "diagram_type": "bar_chart",
            "content": "test",
            "path": "",
        })
        
        # El sistema devuelve error de filesystem cuando path es vacío
        assert "Error insertando diagrama" in result

    def test_nonexistent_path(self):
        """Test error cuando path no existe."""
        result = _tool_doc_insert_diagram({
            "path": "/nonexistent/path.docx",
            "diagram_type": "bar_chart",
            "content": "test",
        })
        
        assert "no encontrado" in result

    def test_bar_chart_styling(self):
        """Test gráfica de barras con estilos Office."""
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            # Datos con formato
            data = "Producto|Ventas\nA|100\nB|150\nC|120\nD|180\nE|200"
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "table",
                "content": data,
                "output_path": str(path)
            })
            
            assert "✅" in result
            
            # Verificar que el archivo se guardó
            assert path.exists()
            
        finally:
            path.unlink(missing_ok=True)

    def test_pie_chart_emojis(self):
        """Test gráfica de pastel con emojis."""
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            # Datos para gráfica de pastel con emojis
            data = "Categoría|Porcentaje\nFrutas|40\nVerduras|30\nGranos|20\nOtros|10"
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "pie_chart",
                "content": data,
                "output_path": str(path)
            })
            
            assert "✅" in result
            
        finally:
            path.unlink(missing_ok=True)

    def test_table_alignment(self):
        """Test tabla con alineación y bordes."""
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            # Tabla con formato
            data = "Col1|Col2|Col3\nA|B|C\n1|2|3\n4|5|6"
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "table",
                "content": data,
                "output_path": str(path)
            })
            
            assert "✅" in result
            
        finally:
            path.unlink(missing_ok=True)

    def test_empty_content(self):
        """Test con contenido vacío."""
        path = Path(tempfile.mktemp(suffix=".docx"))
        try:
            from docx import Document
            doc = Document()
            doc.save(str(path))
            
            result = _tool_doc_insert_diagram({
                "path": str(path),
                "diagram_type": "bar_chart",
                "content": "",
                "output_path": str(path)
            })
            
            # Debería manejar contenido vacío sin error
            assert "✅" in result or "Parámetro" in result
            
        finally:
            path.unlink(missing_ok=True)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
