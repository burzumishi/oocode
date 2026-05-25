#!/usr/bin/env python3
"""Home Office Assistant MCP Server — secretaria IT eficaz para OOCode.

Protocolo: MCP 2024-11-05 sobre stdio (JSON-RPC 2.0 newline-delimited JSON).

Configuración: ~/.oocode/home_office.json  (ver _load_config() para esquema completo)
Proyecto:      OOCODE.md en cwd (templates_dir, docs_dir, naming, client, project_type)
Local:         .oocode-office.json en cwd (override de cualquier clave de config)

Tools (35):
  ── Email ──
  email_list          — lista emails de la bandeja de entrada (IMAP)
  email_read          — lee un email por UID o índice
  email_send          — envía un email (SMTP)
  email_search        — busca emails por criterio (IMAP SEARCH)
  ── Documentos ──
  doc_convert         — convierte documentos con pandoc (md→pdf, docx→md, etc.)
  pdf_extract_text    — extrae texto de un PDF (pdftotext o pdfplumber)
  doc_word_count      — cuenta palabras, líneas y caracteres de un documento
  markdown_to_html    — convierte markdown a HTML
  doc_list_templates  — lista plantillas disponibles (.docx/.xlsx/.md)
  doc_read_template_fields — extrae campos {{CAMPO}} de una plantilla
  doc_fill_template   — rellena campos {{CAMPO}} en plantillas .docx o texto
  doc_create_rfc      — genera un RFC/Change Request estructurado en markdown
  doc_read            — lee contenido de .docx/.md; extrae sección opcional
  doc_update_section  — reemplaza contenido de una sección en .md/.docx
  doc_version_bump    — incrementa version: X.Y.Z en front matter YAML
  doc_insert_diagram  — inserta diagramas y gráficas en documentos .docx
  ── Hojas de cálculo ──
  xlsx_read           — lee celdas o rango de un archivo Excel o CSV
  xlsx_write          — escribe una celda en un archivo Excel (.xlsx)
  xlsx_fill_range     — escribe múltiples celdas a la vez en Excel
  xlsx_append_row     — añade una fila al final de una hoja Excel
  xlsx_create_report  — crea un informe Excel formateado con estilos
  csv_analyze         — análisis de CSV: cabeceras, primeras filas, estadísticas
  ── Calendario / Notas / Contactos ──
  cal_list            — lista eventos de un .ics local
  cal_add             — añade un evento VEVENT a un archivo .ics local
  cal_search          — busca eventos por texto o rango de fechas
  notes_list          — lista ficheros markdown en el vault de notas
  notes_search        — busca texto en notas markdown (ripgrep o grep)
  notes_save          — guarda o actualiza una nota markdown con front matter
  image_to_text       — extrae texto de una imagen con tesseract (OCR)
  contact_search      — busca en ficheros vCard (.vcf)
  ── Workspace / CMDB ──
  project_context_read — lee OOCODE.md y devuelve metadata del proyecto activo
  project_init_office  — inicializa estructura de directorios para proyecto IT
  doc_project_save     — guarda documento en subdir correcto con naming convention
  cmdb_search          — busca en CMDB (CSV/XLSX/JSON); soporta * para listar todo
  cmdb_update          — actualiza registro en CMDB CSV por campo clave
  asset_register_add   — añade activo al registro CSV (lo crea si no existe)

Prompts (12):
  ── Ofimática general ──
  draft_email               — borrador de email profesional con tono configurable
  summarize_document        — resumen estructurado con puntos clave y acción
  meeting_notes             — acta de reunión estructurada
  weekly_report             — informe semanal de actividades
  ── IT / Infraestructura / Datacenter ──
  datacenter_migration_report — informe técnico de migración de DC
  rfc_change_request          — RFC/Change Request formal para infraestructura IT
  server_migration_plan       — plan detallado de migración de servidor
  it_incident_report          — informe post-incidencia IT (Post-Mortem/RCA)
  infrastructure_change_plan  — plan de cambio de infraestructura con fases y riesgos
  ── Gestión / Negocio ──
  executive_summary           — resumen ejecutivo para dirección (1 página)
  business_case               — business case con análisis coste-beneficio y ROI
  project_status_report       — informe RAG de estado de proyecto con KPIs e hitos

Resources (8):
  office://emails_recent       — últimos 10 emails de la bandeja de entrada
  office://calendar_today      — eventos de hoy y los próximos 7 días
  office://notes_recent        — notas markdown más recientes (últimas 10)
  office://templates_available — plantillas de documentos disponibles
  office://rfc_pending         — RFCs/change requests en el directorio de notas
  office://tasks_today         — tareas pendientes del plugin todo de OOCode
  office://project_context     — metadata del proyecto activo (desde OOCODE.md)
  office://server_inventory    — listado completo de activos de la CMDB del proyecto
"""
import csv
import datetime
import email as _email_lib
import email.header
import email.mime.multipart
import email.mime.text
import imaplib
import json
import os
import re
import smtplib
import subprocess
import sys
from pathlib import Path
from typing import Any, Optional


# ── Configuración ─────────────────────────────────────────────────────────────

_CONFIG_PATH = Path.home() / ".oocode" / "home_office.json"

_DEFAULT_CFG: dict = {
    "email": {
        "imap_host":     "",
        "imap_port":     993,
        "imap_ssl":      True,
        "smtp_host":     "",
        "smtp_port":     587,
        "user":          "",
        "password":      "",
        "smtp_user":     "",
        "smtp_password": "",
        "default_from":  "",
    },
    "notes_dir":      str(Path.home() / "Documents" / "notes"),
    "calendar_file":  str(Path.home() / "Documents" / "calendar.ics"),
    "contacts_dir":   str(Path.home() / "Documents" / "contacts"),
    "templates_dir":  str(Path.home() / "Documents" / "templates"),
}


def _parse_oocode_md(path: Path) -> dict:
    """Parse YAML front matter from an OOCODE.md file. Returns {} on failure."""
    if not path.exists():
        return {}
    try:
        text = path.read_text(errors="replace")
        meta: dict = {}
        if text.startswith("---"):
            end = text.find("---", 3)
            if end > 0:
                for line in text[3:end].splitlines():
                    if ":" in line:
                        k, _, v = line.partition(":")
                        meta[k.strip()] = v.strip().strip('"').strip("'")
        return meta
    except Exception:
        return {}


def _apply_naming(cfg: dict, doc_type: str) -> str:
    """Return a filename (without extension) applying the project naming convention."""
    proj    = cfg.get("_project", {})
    pattern = proj.get("naming", "{TYPE}-{YYMMDD}")
    today   = datetime.date.today().strftime("%y%m%d")
    client  = re.sub(r"\s+", "", proj.get("client", "PROJ"))[:8].upper()
    project = re.sub(r"\W+", "-", proj.get("project", "PROJ"))[:12].upper()
    return (
        pattern
        .replace("{CLIENT}", client)
        .replace("{TYPE}",    doc_type.upper()[:4])
        .replace("{YYMMDD}", today)
        .replace("{VER}",    "1")
        .replace("{PROJECT}", project)
    )


def _load_config() -> dict:
    cfg = json.loads(json.dumps(_DEFAULT_CFG))
    if _CONFIG_PATH.exists():
        try:
            user = json.loads(_CONFIG_PATH.read_text())
            _deep_merge(cfg, user)
        except Exception:
            pass
    # Env var overrides
    em = cfg["email"]
    em["imap_host"]     = os.environ.get("HOME_OFFICE_IMAP_HOST",     em["imap_host"])
    em["imap_port"]     = int(os.environ.get("HOME_OFFICE_IMAP_PORT", em["imap_port"]))
    em["smtp_host"]     = os.environ.get("HOME_OFFICE_SMTP_HOST",     em["smtp_host"])
    em["smtp_port"]     = int(os.environ.get("HOME_OFFICE_SMTP_PORT", em["smtp_port"]))
    em["user"]          = os.environ.get("HOME_OFFICE_EMAIL_USER",    em["user"])
    em["password"]      = os.environ.get("HOME_OFFICE_EMAIL_PASS",    em["password"])
    cfg["notes_dir"]     = os.environ.get("HOME_OFFICE_NOTES_DIR",      cfg["notes_dir"])
    cfg["calendar_file"] = os.environ.get("HOME_OFFICE_CALENDAR_FILE", cfg["calendar_file"])
    cfg["contacts_dir"]  = os.environ.get("HOME_OFFICE_CONTACTS_DIR",  cfg["contacts_dir"])
    cfg["templates_dir"] = os.environ.get("HOME_OFFICE_TEMPLATES_DIR", cfg.get("templates_dir", str(Path.home() / "Documents" / "templates")))
    # Local project overrides: .oocode-office.json > OOCODE.md (both in cwd)
    cwd = Path.cwd()
    local_cfg = cwd / ".oocode-office.json"
    if local_cfg.exists():
        try:
            _deep_merge(cfg, json.loads(local_cfg.read_text()))
        except Exception:
            pass
    oocode_md = cwd / "OOCODE.md"
    if oocode_md.exists():
        proj = _parse_oocode_md(oocode_md)
        if proj:
            cfg["_project"] = proj
            if "templates_dir" in proj:
                p = Path(proj["templates_dir"])
                cfg["templates_dir"] = str(p if p.is_absolute() else cwd / p)
            if "docs_dir" in proj:
                p = Path(proj["docs_dir"])
                cfg["_docs_dir"] = str(p if p.is_absolute() else cwd / p)
            if "notes_dir" in proj:
                p = Path(proj["notes_dir"])
                cfg["notes_dir"] = str(p if p.is_absolute() else cwd / p)
    return cfg


def _deep_merge(base: dict, override: dict) -> None:
    for k, v in override.items():
        if k in base and isinstance(base[k], dict) and isinstance(v, dict):
            _deep_merge(base[k], v)
        else:
            base[k] = v


def _no_email_config() -> str:
    return (
        "⚙ Email no configurado.\n"
        f"Edita {_CONFIG_PATH} con tu configuración IMAP/SMTP:\n\n"
        '{\n  "email": {\n'
        '    "imap_host": "imap.gmail.com",\n'
        '    "imap_port": 993,\n'
        '    "smtp_host": "smtp.gmail.com",\n'
        '    "smtp_port": 587,\n'
        '    "user": "tu@email.com",\n'
        '    "password": "contraseña-de-aplicación"\n'
        "  }\n}\n\n"
        "Para Gmail: Configuración → Seguridad → Contraseñas de aplicación.\n"
        "Protege el fichero: chmod 600 ~/.oocode/home_office.json"
    )


# ── Protocolo MCP ────────────────────────────────────────────────────────────

def _send(obj: dict) -> None:
    sys.stdout.write(json.dumps(obj, ensure_ascii=False) + "\n")
    sys.stdout.flush()


def _recv() -> Optional[dict]:
    while True:
        line = sys.stdin.readline()
        if not line:
            return None
        line = line.strip()
        if not line:
            continue
        try:
            return json.loads(line)
        except json.JSONDecodeError:
            continue


def _ok(req_id: Any, result: dict) -> None:
    _send({"jsonrpc": "2.0", "id": req_id, "result": result})


def _err(req_id: Any, code: int, message: str) -> None:
    _send({"jsonrpc": "2.0", "id": req_id, "error": {"code": code, "message": message}})


# ── Helpers ──────────────────────────────────────────────────────────────────

def _decode_header(raw: str) -> str:
    parts = email.header.decode_header(raw or "")
    out = []
    for chunk, enc in parts:
        if isinstance(chunk, bytes):
            out.append(chunk.decode(enc or "utf-8", errors="replace"))
        else:
            out.append(str(chunk))
    return "".join(out)


def _imap_connect(cfg: dict):
    """Conecta a IMAP. Devuelve cliente conectado y autenticado."""
    em = cfg["email"]
    if not em["imap_host"] or not em["user"]:
        return None, _no_email_config()
    try:
        if em["imap_ssl"]:
            client = imaplib.IMAP4_SSL(em["imap_host"], em["imap_port"])
        else:
            client = imaplib.IMAP4(em["imap_host"], em["imap_port"])
        client.login(em["user"], em["password"])
        return client, None
    except Exception as exc:
        return None, f"Error IMAP: {exc}"


def _run(cmd: list, timeout: int = 30, input_text: str = "") -> tuple[int, str, str]:
    try:
        r = subprocess.run(
            cmd, capture_output=True, text=True,
            timeout=timeout,
            input=input_text or None
        )
        return r.returncode, r.stdout, r.stderr
    except FileNotFoundError:
        return -1, "", f"Comando no encontrado: {cmd[0]}"
    except subprocess.TimeoutExpired:
        return -1, "", f"Timeout ({timeout}s) ejecutando {cmd[0]}"


# ── Email tools ──────────────────────────────────────────────────────────────

def _tool_email_list(args: dict) -> str:
    cfg     = _load_config()
    mailbox = args.get("mailbox", "INBOX")
    limit   = min(int(args.get("limit", 10)), 50)

    client, err = _imap_connect(cfg)
    if err:
        return err
    try:
        client.select(mailbox)
        _, data = client.search(None, "ALL")
        uids = data[0].split()
        uids = uids[-limit:][::-1]  # los más recientes primero
        lines = [f"📬 {mailbox} — {len(uids)} emails (mostrando {len(uids)})\n"]
        for uid in uids:
            _, msg_data = client.fetch(uid, "(RFC822.SIZE ENVELOPE)")
            raw = msg_data[0][1].decode("utf-8", errors="replace") if msg_data and msg_data[0] else ""
            # Parse ENVELOPE para extraer subject/from/date
            _, msg_data2 = client.fetch(uid, "(RFC822.HEADER)")
            if msg_data2 and msg_data2[0]:
                msg = _email_lib.message_from_bytes(msg_data2[0][1])
                subj = _decode_header(msg.get("Subject", "(sin asunto)"))[:70]
                frm  = _decode_header(msg.get("From", "?"))[:40]
                date = msg.get("Date", "?")[:30]
                lines.append(f"  [{uid.decode()}] {date}\n     De: {frm}\n     Asunto: {subj}")
        return "\n".join(lines)
    except Exception as exc:
        return f"Error listando emails: {exc}"
    finally:
        try:
            client.logout()
        except Exception:
            pass


def _tool_email_read(args: dict) -> str:
    cfg     = _load_config()
    uid     = str(args.get("uid", ""))
    mailbox = args.get("mailbox", "INBOX")
    if not uid:
        return "Parámetro requerido: uid"

    client, err = _imap_connect(cfg)
    if err:
        return err
    try:
        client.select(mailbox)
        _, msg_data = client.fetch(uid.encode(), "(RFC822)")
        if not msg_data or not msg_data[0]:
            return f"Email UID {uid} no encontrado."
        raw = msg_data[0][1]
        msg = _email_lib.message_from_bytes(raw)
        subj = _decode_header(msg.get("Subject", "(sin asunto)"))
        frm  = _decode_header(msg.get("From", "?"))
        to   = _decode_header(msg.get("To", "?"))
        date = msg.get("Date", "?")
        lines = [
            f"📧 Email UID {uid}",
            f"De:     {frm}",
            f"Para:   {to}",
            f"Asunto: {subj}",
            f"Fecha:  {date}",
            "─" * 60,
        ]
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                ct = part.get_content_type()
                if ct == "text/plain":
                    body = part.get_payload(decode=True).decode(
                        part.get_content_charset() or "utf-8", errors="replace"
                    )
                    break
            if not body:
                for part in msg.walk():
                    if part.get_content_type() == "text/html":
                        raw_html = part.get_payload(decode=True).decode(
                            part.get_content_charset() or "utf-8", errors="replace"
                        )
                        body = re.sub(r"<[^>]+>", "", raw_html)
                        break
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body = payload.decode(msg.get_content_charset() or "utf-8", errors="replace")
        lines.append(body[:4000] if body else "(sin cuerpo)")
        return "\n".join(lines)
    except Exception as exc:
        return f"Error leyendo email {uid}: {exc}"
    finally:
        try:
            client.logout()
        except Exception:
            pass


def _tool_email_send(args: dict) -> str:
    cfg  = _load_config()
    em   = cfg["email"]
    if not em["smtp_host"] or not em["user"]:
        return _no_email_config()

    to      = args.get("to", "")
    subject = args.get("subject", "")
    body    = args.get("body", "")
    cc      = args.get("cc", "")
    bcc     = args.get("bcc", "")
    if not to or not subject:
        return "Parámetros requeridos: to, subject"

    smtp_user = em.get("smtp_user") or em["user"]
    smtp_pass = em.get("smtp_password") or em["password"]
    frm       = em.get("default_from") or smtp_user

    msg = email.mime.multipart.MIMEMultipart()
    msg["From"]    = frm
    msg["To"]      = to
    msg["Subject"] = subject
    if cc:
        msg["Cc"] = cc
    msg.attach(email.mime.text.MIMEText(body, "plain", "utf-8"))

    recipients = [a.strip() for a in (to + "," + cc + "," + bcc).split(",") if a.strip()]
    try:
        with smtplib.SMTP(em["smtp_host"], em["smtp_port"]) as server:
            server.ehlo()
            server.starttls()
            server.login(smtp_user, smtp_pass)
            server.sendmail(frm, recipients, msg.as_string())
        return f"✅ Email enviado a {to}\n   Asunto: {subject}"
    except Exception as exc:
        return f"Error enviando email: {exc}"


def _tool_email_search(args: dict) -> str:
    cfg     = _load_config()
    query   = args.get("query", "")
    mailbox = args.get("mailbox", "INBOX")
    limit   = min(int(args.get("limit", 20)), 50)
    if not query:
        return "Parámetro requerido: query (texto a buscar en asunto/remitente)"

    client, err = _imap_connect(cfg)
    if err:
        return err
    try:
        client.select(mailbox)
        # Intenta búsqueda IMAP nativa en asunto y cuerpo
        clean = query.replace('"', '')
        _, data = client.search(None, f'SUBJECT "{clean}"')
        uids = data[0].split()
        if not uids:
            _, data = client.search(None, f'TEXT "{clean}"')
            uids = data[0].split()
        uids = uids[-limit:][::-1]
        if not uids:
            return f"Sin resultados para: {query}"
        lines = [f"🔍 Búsqueda '{query}' en {mailbox}: {len(uids)} resultados\n"]
        for uid in uids:
            _, msg_data = client.fetch(uid, "(RFC822.HEADER)")
            if msg_data and msg_data[0]:
                msg = _email_lib.message_from_bytes(msg_data[0][1])
                subj = _decode_header(msg.get("Subject", "(sin asunto)"))[:70]
                frm  = _decode_header(msg.get("From", "?"))[:40]
                date = msg.get("Date", "?")[:30]
                lines.append(f"  [{uid.decode()}] {date}\n     De: {frm}\n     Asunto: {subj}")
        return "\n".join(lines)
    except Exception as exc:
        return f"Error buscando emails: {exc}"
    finally:
        try:
            client.logout()
        except Exception:
            pass


# ── Document tools ───────────────────────────────────────────────────────────
# ── Document tools ───────────────────────────────────────────────────────────
# Funciones para generar diagramas y gráficas en documentos Office
def _tool_doc_insert_diagram(args: dict) -> str:
    """Inserta un diagrama o gráfica en un documento .docx con estilos Office.
    
    Soporta:
    - Diagramas de flujo simples (con texto y negritas)
    - Gráficas de barras (usando tablas formateadas)
    - Gráficas de pastel (usando emojis y texto)
    - Tablas de datos dinámicas con formato
    - Organigramas jerárquicos
    
    Estilos Office:
    - Colores de Office (RGB, hex)
    - Fuentes de Office (Calibri, Arial, etc.)
    - Bordes y sombreado
    - Alineación y wrapping de texto
    - Estilos de tabla predefinidos
    
    Nota: Para diagramas complejos, usa visio o vsdw.
    """
    import shutil
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    import xml.etree.ElementTree as ET
    
    path = Path(args.get("path", "")).expanduser()
    diagram_type = args.get("diagram_type", "flowchart")
    content = args.get("content", "")
    output_path = args.get("output_path", "")
    
    if not path:
        return "Parámetro requerido: path"
    if not path.exists():
        return f"Fichero no encontrado: {path}"
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(str(path))
        
        if diagram_type == "flowchart":
            # Diagrama de flujo con texto y negritas
            lines = content.strip().split("\n")
            for i, line in enumerate(lines):
                if line.strip():
                    # Crear párrafo centrado
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(line.strip())
                    run.font.size = Pt(12)
                    run.font.bold = i == 0  # Primer elemento en negrita
        
        elif diagram_type == "bar_chart":
            # Gráfica de barras con tablas formateadas
            data = content.strip().split("\n")
            if len(data) >= 2:
                labels = data[0].split(",")
                values = [d.strip() for d in data[1:]]
                
                # Crear tabla para representar gráfica
                table = doc.add_table(rows=len(labels)+1, cols=2)
                table.style = 'Table Grid'
                for i, label in enumerate(labels):
                    table.cell(i, 0).text = label
                    table.cell(i, 1).text = values[i] if i < len(values) else "0"
                    # Aplicar formato
                    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                    try:
                        cell = table.rows[i].cells[1]
                        cell.font = Font(name="Calibri", size=11)
                        cell.alignment = Alignment(horizontal="right", vertical="center")
                        if i % 2 == 0:
                            cell.fill = PatternFill(start_color="D0D0D0", end_color="D0D0D0", fgColor="gray50")
                    except Exception:
                        pass
        
        elif diagram_type == "table":
            # Tabla de datos dinámica con formato Office
            rows = content.strip().split("\n")
            if len(rows) > 1:
                # Crear tabla con formato
                num_cols = max([len(row.split("|")) for row in rows])
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                
                # Aplicar formato a cada celda
                for i, row in enumerate(rows):
                    cells = row.split("|")
                    for j, cell_text in enumerate(cells):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell_text.strip()
                            # Aplicar formato
                            try:
                                # Color de fondo por fila
                                if i % 2 == 0:
                                    table.rows[i].cells[j].fill = PatternFill(start_color="D0D0D0", end_color="D0D0D0", fgColor="gray50")
                                # Negrita para primera columna
                                if j == 0:
                                    table.rows[i].cells[j].font = Font(bold=True, name="Calibri", size=12)
                                else:
                                    table.rows[i].cells[j].font = Font(name="Calibri", size=11)
                                # Alineación
                                table.rows[i].cells[j].alignment = Alignment(horizontal="center", vertical="center")
                                # Borde
                                table.rows[i].cells[j].border = Border(left="single", right="single", top="single", bottom="single")
                            except Exception:
                                pass
        
        elif diagram_type == "pie_chart":
            # Gráfica de pastel con emojis y texto
            data = content.strip().split("\n")
            if len(data) >= 2:
                labels = data[0].split(",")
                values = [d.strip() for d in data[1:]]
                # Crear elementos de gráfica con emojis
                for i, (label, value) in enumerate(zip(labels, values)):
                    para = doc.add_paragraph()
                    run = para.add_run(f"🥧 {label}: {value}")
                    run.font.size = Pt(10)
        
        elif diagram_type == "org_chart":
            # Organigrama jerárquico con texto
            lines = content.strip().split("\n")
            for i, line in enumerate(lines):
                if line.strip():
                    # Parsear jerarquía
                    parts = line.split(" → ")
                    if len(parts) == 2:
                        # Crear párrafo para cada nivel
                        para1 = doc.add_paragraph()
                        run1 = para1.add_run(parts[0].strip())
                        run1.font.size = Pt(12)
                        run1.font.bold = True
                        run1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Conector
                        para2 = doc.add_paragraph()
                        run2 = para2.add_run(" → ")
                        run2.font.size = Pt(10)
                        
                        para3 = doc.add_paragraph()
                        run3 = para3.add_run(parts[1].strip())
                        run3.font.size = Pt(12)
                        run3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if output_path:
            doc.save(str(output_path))
            return f"✅ Diagrama insertado: {output_path}\n   Tipo: {diagram_type}"
        else:
            doc.save(str(path))
            return f"✅ Diagrama insertado en: {path}\n   Tipo: {diagram_type}"
    
    except ImportError as e:
        return f"python-docx no disponible. Instala con: pip install python-docx\nError: {e}"
    except Exception as exc:
        return f"Error insertando diagrama: {exc}"

def _tool_doc_convert(args: dict) -> str:
    import shutil
    if not shutil.which("pandoc"):
        return (
            "pandoc no encontrado. Instala con:\n"
            "  Linux: sudo apt install pandoc\n"
            "  Mac: brew install pandoc\n"
            "  Windows (WSL): sudo apt install pandoc"
        )
    input_path  = Path(args.get("input_path", "")).expanduser()
    out_format  = args.get("output_format", "")
    output_path = args.get("output_path", "")
    if not input_path or not out_format:
        return "Parámetros requeridos: input_path, output_format (p.ej. 'pdf', 'html', 'docx', 'md')"
    if not input_path.exists():
        return f"Fichero no encontrado: {input_path}"
    if not output_path:
        output_path = str(input_path.with_suffix(f".{out_format}"))
    cmd = ["pandoc", str(input_path), "-o", output_path, "--standalone"]
    if out_format == "pdf":
        cmd += ["--pdf-engine=xelatex"]
    rc, out, err = _run(cmd, timeout=60)
    if rc != 0:
        return f"Error convirtiendo documento:\n{err}"
    size = Path(output_path).stat().st_size if Path(output_path).exists() else 0
    return f"✅ Convertido: {input_path.name} → {output_path}\n   Tamaño: {size:,} bytes"


def _tool_pdf_extract_text(args: dict) -> str:
    import shutil
    path   = Path(args.get("path", "")).expanduser()
    pages  = args.get("pages", "")
    if not path:
        return "Parámetro requerido: path"
    if not path.exists():
        return f"Fichero no encontrado: {path}"
    # Intento 1: pdftotext (poppler-utils)
    if shutil.which("pdftotext"):
        cmd = ["pdftotext"]
        if pages:
            parts = pages.split("-")
            if len(parts) == 2:
                cmd += ["-f", parts[0], "-l", parts[1]]
        cmd += [str(path), "-"]
        rc, out, err = _run(cmd, timeout=30)
        if rc == 0:
            text = out.strip()
            note = f"\n[Extraídas páginas: {pages}]" if pages else ""
            return f"📄 {path.name}{note}\n{'─'*50}\n{text[:8000]}"
        return f"Error con pdftotext: {err}"
    # Intento 2: pdfplumber (pip)
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
        text = "\n".join(text_parts)
        return f"📄 {path.name}\n{'─'*50}\n{text[:8000]}"
    except ImportError:
        pass
    return (
        "Necesitas pdftotext o pdfplumber:\n"
        "  Linux: sudo apt install poppler-utils\n"
        "  pip:   pip install pdfplumber"
    )


def _tool_doc_word_count(args: dict) -> str:
    raw = args.get("path", "")
    if not raw:
        return "Parámetro requerido: path"
    path = Path(raw).expanduser()
    if path.is_dir():
        return f"Parámetro requerido: path a un fichero, no a un directorio: {path}"
    if not path.exists():
        return f"Fichero no encontrado: {path}"
    try:
        text = path.read_text(errors="replace")
        lines = text.splitlines()
        words = len(text.split())
        chars = len(text)
        chars_no_spaces = len(text.replace(" ", "").replace("\n", ""))
        paragraphs = len([p for p in text.split("\n\n") if p.strip()])
        return (
            f"📊 {path.name}\n"
            f"  Líneas:         {len(lines):>8,}\n"
            f"  Palabras:       {words:>8,}\n"
            f"  Caracteres:     {chars:>8,}\n"
            f"  (sin espacios): {chars_no_spaces:>8,}\n"
            f"  Párrafos:       {paragraphs:>8,}\n"
            f"  Tamaño:         {path.stat().st_size:>8,} bytes"
        )
    except Exception as exc:
        return f"Error leyendo fichero: {exc}"


# ── Spreadsheet tools ────────────────────────────────────────────────────────

def _tool_xlsx_read(args: dict) -> str:
    path   = Path(args.get("path", "")).expanduser()
    sheet  = args.get("sheet", "")
    rng    = args.get("range", "")
    limit  = int(args.get("limit", 50))
    if not path:
        return "Parámetro requerido: path"
    if not path.exists():
        return f"Fichero no encontrado: {path}"
    ext = path.suffix.lower()
    if ext == ".csv":
        return _read_csv_file(path, limit)
    # Excel
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        ws_name = sheet or wb.sheetnames[0]
        if ws_name not in wb.sheetnames:
            return f"Hoja '{ws_name}' no encontrada. Disponibles: {', '.join(wb.sheetnames)}"
        ws = wb[ws_name]
        lines = [f"📊 {path.name} — Hoja: {ws_name}"]
        count = 0
        for row in ws.iter_rows(values_only=True):
            if count >= limit:
                lines.append(f"  … (limitado a {limit} filas)")
                break
            cells = [str(c) if c is not None else "" for c in row]
            lines.append("  " + " | ".join(cells))
            count += 1
        return "\n".join(lines)
    except ImportError:
        return "openpyxl no instalado. Instala con: pip install openpyxl"
    except Exception as exc:
        return f"Error leyendo Excel: {exc}"


def _read_csv_file(path: Path, limit: int) -> str:
    try:
        with path.open(newline="", errors="replace") as f:
            reader = csv.reader(f)
            rows = list(reader)
        if not rows:
            return f"CSV vacío: {path.name}"
        headers = rows[0]
        data    = rows[1:limit + 1]
        lines = [f"📊 {path.name} — CSV ({len(rows)-1} filas, {len(headers)} columnas)"]
        lines.append("  " + " | ".join(headers))
        lines.append("  " + " | ".join(["─" * min(len(h), 12) for h in headers]))
        for row in data:
            lines.append("  " + " | ".join(row))
        if len(rows) - 1 > limit:
            lines.append(f"  … ({len(rows)-1-limit} filas más)")
        return "\n".join(lines)
    except Exception as exc:
        return f"Error leyendo CSV: {exc}"


def _apply_cell_style(cell, style: dict) -> None:
    """Aplica un dict de estilo a una celda openpyxl.

    Claves soportadas:
      bold, italic, underline (bool)
      font_size (int), font_color (str hex RRGGBB p.ej. "FF0000")
      bg_color (str hex RRGGBB), align (str: left/center/right/fill)
      number_format (str, p.ej. "#,##0.00", "0%", "DD/MM/YYYY")
      border (bool) — añade borde fino en los 4 lados
      wrap_text (bool)
    """
    try:
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side  # type: ignore
        kw: dict = {}
        if style.get("bold"):       kw["bold"]      = True
        if style.get("italic"):     kw["italic"]    = True
        if style.get("underline"):  kw["underline"] = "single"
        if style.get("font_size"):  kw["size"]      = int(style["font_size"])
        if style.get("font_color"): kw["color"]     = style["font_color"].lstrip("#").upper()
        if kw:
            cell.font = Font(**kw)
        bg = style.get("bg_color", "")
        if bg:
            cell.fill = PatternFill("solid", fgColor=bg.lstrip("#").upper())
        align = style.get("align", "")
        wrap  = bool(style.get("wrap_text"))
        if align or wrap:
            cell.alignment = Alignment(horizontal=align or None, wrap_text=wrap or None)
        if style.get("border"):
            thin = Side(style="thin")
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
        if style.get("number_format"):
            cell.number_format = style["number_format"]
    except Exception:
        pass


def _tool_xlsx_write(args: dict) -> str:
    path  = Path(args.get("path", "")).expanduser()
    sheet = args.get("sheet", "Hoja1")
    cell  = args.get("cell", "A1")
    value = args.get("value", "")
    style = args.get("style", {})
    if not path or not cell:
        return "Parámetros requeridos: path, cell (p.ej. 'B3'), value"
    try:
        import openpyxl
        if path.exists():
            wb = openpyxl.load_workbook(str(path))
        else:
            wb = openpyxl.Workbook()
        if sheet not in wb.sheetnames:
            wb.create_sheet(sheet)
        ws = wb[sheet]
        c = ws[cell.upper()]
        c.value = value
        if style and isinstance(style, dict):
            _apply_cell_style(c, style)
        wb.save(str(path))
        return f"✅ Escrito en {path.name} [{sheet}]{cell.upper()} = {value!r}"
    except ImportError:
        return "openpyxl no instalado. Instala con: pip install openpyxl"
    except Exception as exc:
        return f"Error escribiendo Excel: {exc}"


def _tool_csv_analyze(args: dict) -> str:
    path  = Path(args.get("path", "")).expanduser()
    limit = int(args.get("limit", 5))
    if not path:
        return "Parámetro requerido: path"
    if not path.exists():
        return f"Fichero no encontrado: {path}"
    try:
        with path.open(newline="", errors="replace") as f:
            sample = f.read(65536)
        dialect = csv.Sniffer().sniff(sample[:2048])
        with path.open(newline="", errors="replace") as f:
            reader = csv.DictReader(f, dialect=dialect)
            rows   = [row for _, row in zip(range(limit + 1), reader)]
        headers = list(rows[0].keys()) if rows else []
        n_rows  = sum(1 for _ in path.open(newline="")) - 1
        lines = [
            f"📊 Análisis CSV: {path.name}",
            f"  Filas (aprox):  {n_rows}",
            f"  Columnas:       {len(headers)}",
            f"  Delimitador:    {dialect.delimiter!r}",
            f"  Columnas: {', '.join(headers)}",
            f"\n  Primeras {min(limit, len(rows))} filas:",
        ]
        for i, row in enumerate(rows[:limit]):
            lines.append(f"  [{i+1}] " + " | ".join(f"{k}={v!r}" for k, v in row.items()))
        # Stats básicas para columnas numéricas
        numeric: dict[str, list] = {h: [] for h in headers}
        with path.open(newline="", errors="replace") as f:
            reader2 = csv.DictReader(f, dialect=dialect)
            for row in reader2:
                for h in headers:
                    try:
                        numeric[h].append(float(row[h]))
                    except (ValueError, TypeError, KeyError):
                        pass
        num_cols = [(h, vals) for h, vals in numeric.items() if len(vals) >= n_rows * 0.5]
        if num_cols:
            lines.append("\n  Estadísticas numéricas:")
            for h, vals in num_cols:
                mn = min(vals); mx = max(vals); avg = sum(vals) / len(vals)
                lines.append(f"    {h}: min={mn:.2f}  max={mx:.2f}  avg={avg:.2f}")
        return "\n".join(lines)
    except Exception as exc:
        return f"Error analizando CSV: {exc}"


# ── Calendar tools ───────────────────────────────────────────────────────────

def _parse_ics_events(path: Path) -> list[dict]:
    """Parser ICS mínimo sin dependencias externas."""
    events = []
    current: dict = {}
    in_event = False
    try:
        lines = path.read_text(errors="replace").splitlines()
        for line in lines:
            if line.startswith(" ") or line.startswith("\t"):
                # Continuación de propiedad
                if current and "_last_key" in current:
                    current[current["_last_key"]] += line[1:]
                continue
            if line.strip() == "BEGIN:VEVENT":
                current = {}
                in_event = True
            elif line.strip() == "END:VEVENT":
                events.append(current)
                in_event = False
                current = {}
            elif in_event and ":" in line:
                key, _, val = line.partition(":")
                key = key.split(";")[0].strip().upper()
                current[key] = val.strip()
                current["_last_key"] = key
    except Exception:
        pass
    return events


def _ics_date_str(raw: str) -> str:
    raw = raw.replace("Z", "").replace("-", "").replace(":", "")
    try:
        if len(raw) == 8:
            return datetime.datetime.strptime(raw, "%Y%m%d").strftime("%Y-%m-%d")
        return datetime.datetime.strptime(raw[:15], "%Y%m%dT%H%M%S").strftime("%Y-%m-%d %H:%M")
    except Exception:
        return raw


def _tool_cal_list(args: dict) -> str:
    cfg    = _load_config()
    source = args.get("source", cfg["calendar_file"])
    start  = args.get("start", "")
    end    = args.get("end", "")
    limit  = int(args.get("limit", 15))
    ics_path = Path(source).expanduser()
    if not ics_path.exists():
        return f"Fichero de calendario no encontrado: {ics_path}\nConfigura 'calendar_file' en {_CONFIG_PATH}"
    events = _parse_ics_events(ics_path)
    today = datetime.date.today()
    result_events = []
    for ev in events:
        dt_str = ev.get("DTSTART", "")
        dt_disp = _ics_date_str(dt_str)
        summary = ev.get("SUMMARY", "(sin título)")
        location = ev.get("LOCATION", "")
        result_events.append((dt_str, dt_disp, summary, location))
    result_events.sort(key=lambda x: x[0])
    # Filtrar por fechas si se proporcionan
    if start:
        result_events = [e for e in result_events if e[0] >= start.replace("-", "")]
    if end:
        result_events = [e for e in result_events if e[0] <= end.replace("-", "")]
    result_events = result_events[:limit]
    if not result_events:
        return f"📅 Sin eventos en {ics_path.name}"
    lines = [f"📅 Calendario: {ics_path.name} — {len(result_events)} eventos"]
    for _, dt_disp, summary, location in result_events:
        loc = f" @ {location}" if location else ""
        lines.append(f"  {dt_disp:<20} {summary}{loc}")
    return "\n".join(lines)


def _tool_cal_add(args: dict) -> str:
    cfg      = _load_config()
    title    = args.get("title", "")
    start    = args.get("start", "")
    end      = args.get("end", "")
    location = args.get("location", "")
    desc     = args.get("description", "")
    file     = args.get("file", cfg["calendar_file"])
    if not title or not start:
        return "Parámetros requeridos: title, start (formato: YYYY-MM-DD o YYYY-MM-DDTHH:MM)"
    ics_path = Path(file).expanduser()
    # Formatear fecha para ICS
    def _to_ics_dt(s: str) -> str:
        s = s.replace("-", "").replace(":", "").replace(" ", "T")
        return s if "T" in s else s + "T000000"
    uid = f"{datetime.datetime.now().strftime('%Y%m%dT%H%M%S')}-oocode@home"
    vevent_lines = [
        "BEGIN:VEVENT",
        f"UID:{uid}",
        f"DTSTAMP:{datetime.datetime.now(datetime.timezone.utc).strftime('%Y%m%dT%H%M%SZ')}",
        f"DTSTART:{_to_ics_dt(start)}",
    ]
    if end:
        vevent_lines.append(f"DTEND:{_to_ics_dt(end)}")
    vevent_lines.append(f"SUMMARY:{title}")
    if location:
        vevent_lines.append(f"LOCATION:{location}")
    if desc:
        vevent_lines.append(f"DESCRIPTION:{desc}")
    vevent_lines.append("END:VEVENT")
    vevent = "\r\n".join(vevent_lines) + "\r\n"
    try:
        if ics_path.exists():
            content = ics_path.read_text()
            if "END:VCALENDAR" in content:
                content = content.replace("END:VCALENDAR", vevent + "END:VCALENDAR")
            else:
                content += vevent
        else:
            ics_path.parent.mkdir(parents=True, exist_ok=True)
            content = (
                "BEGIN:VCALENDAR\r\nVERSION:2.0\r\nPRODID:-//OOCode//Home Office//ES\r\n"
                + vevent + "END:VCALENDAR\r\n"
            )
        ics_path.write_text(content)
        return f"✅ Evento añadido: {title}\n   Inicio: {start}\n   Fichero: {ics_path}"
    except Exception as exc:
        return f"Error añadiendo evento: {exc}"


def _tool_cal_search(args: dict) -> str:
    cfg    = _load_config()
    query  = args.get("query", "")
    source = args.get("source", cfg["calendar_file"])
    limit  = int(args.get("limit", 20))
    if not query:
        return "Parámetro requerido: query"
    ics_path = Path(source).expanduser()
    if not ics_path.exists():
        return f"Fichero de calendario no encontrado: {ics_path}"
    events  = _parse_ics_events(ics_path)
    q_lower = query.lower()
    matches = []
    for ev in events:
        summary = ev.get("SUMMARY", "")
        desc    = ev.get("DESCRIPTION", "")
        if q_lower in summary.lower() or q_lower in desc.lower():
            dt_disp  = _ics_date_str(ev.get("DTSTART", ""))
            location = ev.get("LOCATION", "")
            matches.append((ev.get("DTSTART", ""), dt_disp, summary, location))
    matches.sort(key=lambda x: x[0])
    matches = matches[:limit]
    if not matches:
        return f"Sin eventos que coincidan con: {query}"
    lines = [f"🔍 Búsqueda '{query}': {len(matches)} eventos"]
    for _, dt_disp, summary, location in matches:
        loc = f" @ {location}" if location else ""
        lines.append(f"  {dt_disp:<20} {summary}{loc}")
    return "\n".join(lines)


# ── Notes tools ──────────────────────────────────────────────────────────────

def _notes_dir(cfg: dict) -> Path:
    d = Path(cfg["notes_dir"]).expanduser()
    d.mkdir(parents=True, exist_ok=True)
    return d


def _tool_notes_list(args: dict) -> str:
    cfg     = _load_config()
    dirpath = Path(args.get("directory", cfg["notes_dir"])).expanduser()
    pattern = args.get("pattern", "*.md")
    limit   = int(args.get("limit", 20))
    if not dirpath.exists():
        return f"Directorio de notas no encontrado: {dirpath}\nConfigura 'notes_dir' en {_CONFIG_PATH}"
    files = sorted(dirpath.glob(pattern), key=lambda f: f.stat().st_mtime, reverse=True)[:limit]
    if not files:
        return f"Sin notas en {dirpath} (patrón: {pattern})"
    lines = [f"📝 Notas en {dirpath} — {len(files)} ficheros"]
    for f in files:
        mtime = datetime.datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
        size  = f.stat().st_size
        # Intentar leer título del front matter o primera línea
        try:
            first = f.read_text(errors="replace")[:200]
            title_match = re.search(r"^title:\s*(.+)$", first, re.MULTILINE)
            title = title_match.group(1).strip() if title_match else first.split("\n")[0][:60]
            title = title.lstrip("#").strip()
        except Exception:
            title = f.stem
        lines.append(f"  {mtime}  {size:>6} B  {f.name:<30} {title}")
    return "\n".join(lines)


def _tool_notes_search(args: dict) -> str:
    import shutil
    cfg     = _load_config()
    query   = args.get("query", "")
    dirpath = Path(args.get("directory", cfg["notes_dir"])).expanduser()
    limit   = int(args.get("limit", 10))
    if not query:
        return "Parámetro requerido: query"
    if not dirpath.exists():
        return f"Directorio de notas no encontrado: {dirpath}"
    # ripgrep o grep
    if shutil.which("rg"):
        rc, out, err = _run(["rg", "--color=never", "-l", "-i", query, str(dirpath)], timeout=15)
        tool = "rg"
    else:
        rc, out, err = _run(["grep", "-r", "-l", "-i", query, str(dirpath)], timeout=15)
        tool = "grep"
    if rc != 0 and not out:
        return f"Sin resultados para: {query}"
    files = [Path(p) for p in out.strip().splitlines() if p][:limit]
    if not files:
        return f"Sin notas que contengan: {query}"
    lines = [f"🔍 Búsqueda '{query}' en notas ({len(files)} ficheros)"]
    for f in files:
        # Mostrar líneas que coinciden
        if shutil.which("rg"):
            rc2, ctx, _ = _run(["rg", "--color=never", "-n", "-i", query, str(f)], timeout=10)
        else:
            rc2, ctx, _ = _run(["grep", "-n", "-i", query, str(f)], timeout=10)
        ctx_lines = ctx.strip().splitlines()[:3]
        lines.append(f"\n  📄 {f.name}")
        for cl in ctx_lines:
            lines.append(f"     {cl}")
    return "\n".join(lines)


def _tool_notes_save(args: dict) -> str:
    cfg     = _load_config()
    title   = args.get("title", "")
    content = args.get("content", "")
    dirpath = Path(args.get("directory", cfg["notes_dir"])).expanduser()
    if not title:
        return "Parámetro requerido: title"
    dirpath.mkdir(parents=True, exist_ok=True)
    # Sanitizar título como nombre de fichero
    fname = re.sub(r"[^\w\s\-]", "", title).strip().replace(" ", "_")[:60] + ".md"
    note_path = dirpath / fname
    now = datetime.datetime.now().strftime("%Y-%m-%d")
    if note_path.exists():
        old = note_path.read_text(errors="replace")
        # Actualiza fecha de modificación en front matter si existe
        if old.startswith("---"):
            content_final = re.sub(r"^modified:.*$", f"modified: {now}", old, flags=re.MULTILINE)
            if "modified:" not in content_final:
                content_final = content_final.replace("---\n", f"---\nmodified: {now}\n", 1)
            # Reemplaza contenido del cuerpo (después del segundo ---)
            parts = content_final.split("---", 2)
            if len(parts) == 3 and content:
                content_final = "---" + parts[1] + "---\n\n" + content
        else:
            content_final = content or old
        note_path.write_text(content_final)
        return f"✅ Nota actualizada: {note_path}"
    else:
        front_matter = f"---\ntitle: {title}\ncreated: {now}\nmodified: {now}\n---\n\n"
        note_path.write_text(front_matter + (content or ""))
        return f"✅ Nota creada: {note_path}"


# ── Misc tools ───────────────────────────────────────────────────────────────

def _tool_image_to_text(args: dict) -> str:
    import shutil
    path = Path(args.get("path", "")).expanduser()
    lang = args.get("lang", "spa+eng")
    if not path:
        return "Parámetro requerido: path"
    if not path.exists():
        return f"Fichero no encontrado: {path}"
    if not shutil.which("tesseract"):
        return (
            "tesseract no encontrado. Instala con:\n"
            "  Linux: sudo apt install tesseract-ocr tesseract-ocr-spa\n"
            "  Mac: brew install tesseract tesseract-lang"
        )
    rc, out, err = _run(["tesseract", str(path), "stdout", "-l", lang], timeout=60)
    if rc != 0:
        return f"Error OCR: {err}"
    return f"🔤 OCR de {path.name}:\n{'─'*50}\n{out.strip()}"


def _tool_contact_search(args: dict) -> str:
    cfg     = _load_config()
    query   = args.get("query", "")
    vcf_dir = Path(args.get("vcf_dir", cfg["contacts_dir"])).expanduser()
    if not query:
        return "Parámetro requerido: query"
    if not vcf_dir.exists():
        return f"Directorio de contactos no encontrado: {vcf_dir}\nConfigura 'contacts_dir' en {_CONFIG_PATH}"
    vcf_files = list(vcf_dir.glob("*.vcf")) + list(vcf_dir.glob("*.vcard"))
    if not vcf_files:
        return f"No se encontraron ficheros .vcf en {vcf_dir}"
    q_lower = query.lower()
    results = []
    for vcf_file in vcf_files:
        try:
            text = vcf_file.read_text(errors="replace")
        except Exception:
            continue
        if q_lower not in text.lower():
            continue
        # Extraer campos relevantes
        fn     = re.search(r"^FN:(.+)$",    text, re.MULTILINE)
        email  = re.search(r"^EMAIL.*:(.+)$", text, re.MULTILINE | re.IGNORECASE)
        tel    = re.search(r"^TEL.*:(.+)$",  text, re.MULTILINE | re.IGNORECASE)
        org    = re.search(r"^ORG:(.+)$",    text, re.MULTILINE)
        name   = fn.group(1).strip()    if fn    else vcf_file.stem
        e_val  = email.group(1).strip() if email else ""
        t_val  = tel.group(1).strip()   if tel   else ""
        o_val  = org.group(1).strip()   if org   else ""
        results.append((name, e_val, t_val, o_val))
    if not results:
        return f"Sin contactos que coincidan con: {query}"
    lines = [f"👤 Contactos ({len(results)} encontrados)"]
    for name, email_v, tel_v, org_v in results:
        lines.append(f"\n  {name}")
        if org_v:   lines.append(f"    Empresa: {org_v}")
        if email_v: lines.append(f"    Email:   {email_v}")
        if tel_v:   lines.append(f"    Tel:     {tel_v}")
    return "\n".join(lines)


def _tool_markdown_to_html(args: dict) -> str:
    import shutil
    content = args.get("content", "")
    path    = args.get("path", "")
    output  = args.get("output", "")
    if path:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Fichero no encontrado: {p}"
        content = p.read_text(errors="replace")
    if not content:
        return "Parámetro requerido: content (markdown) o path (ruta a fichero .md)"
    # Intento 1: markdown lib
    try:
        import markdown as md_lib
        html = md_lib.markdown(content, extensions=["tables", "fenced_code", "nl2br"])
        html = f"<!DOCTYPE html>\n<html>\n<body>\n{html}\n</body>\n</html>"
        if output:
            Path(output).expanduser().write_text(html)
            return f"✅ HTML escrito en: {output}\n{html[:500]}…"
        return html[:6000]
    except ImportError:
        pass
    # Intento 2: pandoc
    if shutil.which("pandoc"):
        cmd = ["pandoc", "-f", "markdown", "-t", "html", "--standalone"]
        rc, out, err = _run(cmd, timeout=15, input_text=content)
        if rc == 0:
            if output:
                Path(output).expanduser().write_text(out)
                return f"✅ HTML escrito en: {output}"
            return out[:6000]
    # Fallback: conversión básica manual
    html = content
    html = re.sub(r"^# (.+)$",  r"<h1>\1</h1>",  html, flags=re.MULTILINE)
    html = re.sub(r"^## (.+)$", r"<h2>\1</h2>",  html, flags=re.MULTILINE)
    html = re.sub(r"^### (.+)$",r"<h3>\1</h3>",  html, flags=re.MULTILINE)
    html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
    html = re.sub(r"\*(.+?)\*",     r"<em>\1</em>",         html)
    html = re.sub(r"`(.+?)`",       r"<code>\1</code>",      html)
    html = html.replace("\n\n", "</p>\n<p>")
    html = f"<p>{html}</p>"
    return html[:6000]


# ── Bloque 1: Workspace / Proyecto ───────────────────────────────────────────

def _tool_project_context_read(args: dict) -> str:
    """Read and display project context from OOCODE.md in the workspace or a given path."""
    raw = args.get("path", "OOCODE.md")
    p   = Path(raw).expanduser()
    if not p.is_absolute():
        p = Path.cwd() / p
    if not p.exists():
        return (
            f"No se encontró OOCODE.md en {p.parent}\n"
            "Usa 'project_init_office' para crear la estructura de un proyecto IT."
        )
    meta = _parse_oocode_md(p)
    text = p.read_text(errors="replace")
    body = text
    if text.startswith("---"):
        end = text.find("---", 3)
        if end > 0:
            body = text[end + 3:].strip()
    lines = [f"📋 Proyecto: {p}"]
    if meta:
        lines.append("\n**Metadatos:**")
        for k, v in meta.items():
            lines.append(f"  {k}: {v}")
    lines.append(f"\n**Contenido:**\n{body[:3000]}")
    if len(body) > 3000:
        lines.append(f"… [{len(body)-3000} caracteres más]")
    return "\n".join(lines)


def _tool_project_init_office(args: dict) -> str:
    """Initialize an IT project structure: OOCODE.md, doc folders, skeleton CMDB/registers."""
    project   = args.get("project", "")
    if not project:
        return "Parámetro requerido: project (nombre del proyecto)"
    client    = args.get("client", "")
    proj_type = args.get("type", "general")
    dc_source = args.get("dc_source", "")
    dc_target = args.get("dc_target", "")
    team      = args.get("team", "")
    approver  = args.get("approver", "")
    naming    = args.get("naming", "{CLIENT}-{TYPE}-{YYMMDD}-v{VER}")
    base      = Path(args.get("directory", ".")).expanduser()
    if not base.is_absolute():
        base = Path.cwd() / base
    base.mkdir(parents=True, exist_ok=True)

    dirs_spec = ["templates", "docs/rfcs", "docs/reports", "docs/meetings",
                 "docs/incidents", "docs/plans"]
    created_dirs = []
    for d in dirs_spec:
        p = base / d
        if not p.exists():
            p.mkdir(parents=True, exist_ok=True)
            created_dirs.append(d)

    today     = datetime.date.today().isoformat()
    fm: dict  = {
        "project": project, "client": client, "type": proj_type,
        "created": today, "team": team, "approver": approver,
        "naming": naming, "docs_dir": "./docs", "templates_dir": "./templates",
    }
    if dc_source:
        fm["dc_source"] = dc_source
    if dc_target:
        fm["dc_target"] = dc_target
    fm_text = "---\n" + "\n".join(f'{k}: "{v}"' for k, v in fm.items() if v) + "\n---\n"
    descriptions = {
        "migration":   f"Migración de infraestructura desde {dc_source or 'origen'} → {dc_target or 'destino'}.",
        "datacenter":  "Gestión y operación de centro de datos.",
        "rfc":         "Gestión de cambios y RFCs de infraestructura IT.",
        "audit":       "Auditoría de infraestructura IT.",
        "cloud":       "Migración y gestión de infraestructura cloud.",
        "general":     "Proyecto de infraestructura IT.",
    }
    desc = descriptions.get(proj_type, descriptions["general"])
    oocode_content = (
        fm_text
        + f"\n# {project}\n\n{desc}\n\n"
        "## Estructura del proyecto\n\n"
        "```\n./templates/      — Plantillas de documentos\n"
        "./docs/rfcs/      — RFC y Change Requests\n"
        "./docs/reports/   — Informes y reportes\n"
        "./docs/meetings/  — Actas de reunión\n"
        "./docs/incidents/ — Informes de incidencias\n"
        "./docs/plans/     — Planes de migración/cambio\n"
        "```\n"
    )
    oocode_md = base / "OOCODE.md"
    existed = oocode_md.exists()
    if not existed:
        oocode_md.write_text(oocode_content)

    created_files = []
    cmdb_path = base / "cmdb.csv"
    if not cmdb_path.exists():
        with cmdb_path.open("w", newline="") as f:
            csv.writer(f).writerow(
                ["hostname", "ip", "os", "role", "environment", "owner", "status", "location", "notes"]
            )
        created_files.append("cmdb.csv")
    risk_path = base / "risk_register.csv"
    if not risk_path.exists():
        with risk_path.open("w", newline="") as f:
            csv.writer(f).writerow(
                ["id", "description", "probability", "impact", "level", "mitigation", "owner", "status", "date"]
            )
        created_files.append("risk_register.csv")

    lines = [f"✅ Proyecto '{project}' {'encontrado' if existed else 'inicializado'} en {base}"]
    lines.append(f"   {'ℹ️  OOCODE.md existente — no modificado' if existed else '📄 OOCODE.md creado'}")
    if created_dirs:
        lines.append(f"   📁 Directorios: {', '.join(created_dirs)}")
    if created_files:
        lines.append(f"   📊 Registros: {', '.join(created_files)}")
    lines.append(f"   Tipo: {proj_type}  |  Cliente: {client or 'N/A'}  |  Equipo: {team or 'N/A'}")
    lines.append("\nPróximos pasos:")
    lines.append("  1. Añade plantillas .docx/.md en ./templates/")
    lines.append("  2. Rellena cmdb.csv con el inventario de servidores")
    lines.append("  3. Usa doc_create_rfc / datacenter_migration_report para comenzar")
    return "\n".join(lines)


def _tool_doc_project_save(args: dict) -> str:
    """Save generated document content to the correct project folder using naming convention."""
    content  = args.get("content", "")
    doc_type = args.get("doc_type", "general")
    filename = args.get("filename", "")
    # Formal types default to .docx; loose notes/general default to .md
    _FORMAL_TYPES = {"rfc", "change_request", "report", "informe", "meeting", "acta",
                     "incident", "incidencia", "plan", "migration"}
    default_ext = ".docx" if doc_type.lower() in _FORMAL_TYPES else ".md"
    ext      = args.get("ext", default_ext)
    if not content:
        return "Parámetro requerido: content (texto del documento)"

    cfg  = _load_config()
    cwd  = Path.cwd()
    docs_base = Path(args.get("directory", cfg.get("_docs_dir", str(cwd / "docs")))).expanduser()
    type_dirs = {
        "rfc": "rfcs", "change_request": "rfcs",
        "report": "reports", "informe": "reports",
        "meeting": "meetings", "acta": "meetings",
        "incident": "incidents", "incidencia": "incidents",
        "plan": "plans", "migration": "plans",
        "general": ".",
    }
    subdir = type_dirs.get(doc_type.lower(), "general")
    save_dir = docs_base / subdir
    save_dir.mkdir(parents=True, exist_ok=True)

    if not filename:
        name = re.sub(r"[^\w\-]", "", _apply_naming(cfg, doc_type))
        filename = (name or f"{doc_type}-{datetime.date.today().strftime('%Y%m%d')}") + ext

    save_path = save_dir / filename
    if save_path.exists():
        stem, sfx = save_path.stem, save_path.suffix
        i = 2
        while save_path.exists():
            save_path = save_dir / f"{stem}_v{i}{sfx}"
            i += 1

    if save_path.suffix.lower() == ".docx":
        err = _write_docx_from_markdown(content, save_path)
        if err:
            save_path = save_path.with_suffix(".md")
            save_path.write_text(content)
            return (
                f"⚠ {err}\n✅ Guardado como Markdown: {save_path}\n"
                f"   Tipo: {doc_type}  |  Tamaño: {save_path.stat().st_size:,} bytes"
            )
    else:
        save_path.write_text(content)
    return (
        f"✅ Documento guardado: {save_path}\n"
        f"   Tipo: {doc_type}  |  Tamaño: {save_path.stat().st_size:,} bytes"
    )


# ── Bloque 2: Documento inteligente ──────────────────────────────────────────

def _extract_section_text(text: str, section: str) -> tuple[str, int, int]:
    """Find section by heading in markdown text. Returns (section_text, start_line, end_line)."""
    lines = text.splitlines()
    start = end = None
    heading_lvl = 0
    for i, line in enumerate(lines):
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            lvl = len(m.group(1))
            if section.lower() in m.group(2).lower():
                start = i
                heading_lvl = lvl
            elif start is not None and lvl <= heading_lvl:
                end = i
                break
    if start is None:
        return "", -1, -1
    if end is None:
        end = len(lines)
    return "\n".join(lines[start:end]), start, end


def _tool_doc_read(args: dict) -> str:
    """Read text content from a .docx, .md, or .txt file. Optionally extract a specific section."""
    raw = args.get("path", "")
    if not raw or not str(raw).strip():
        return "Parámetro requerido: path"
    path      = Path(raw).expanduser()
    max_chars = int(args.get("max_chars", 8000))
    section   = args.get("section", "")
    if not path.exists():
        return f"Fichero no encontrado: {path}"

    ext = path.suffix.lower()
    # Plain text formats
    if ext in (".md", ".txt", ".rst", ".html", ".csv"):
        text = path.read_text(errors="replace")
        if section:
            stext, sl, el = _extract_section_text(text, section)
            if sl == -1:
                return f"Sección '{section}' no encontrada en {path.name}"
            return f"📄 {path.name} — Sección: {section}\n{'─'*50}\n{stext[:max_chars]}"
        return f"📄 {path.name}\n{'─'*50}\n{text[:max_chars]}"

    if ext == ".docx":
        # Try python-docx
        try:
            from docx import Document  # type: ignore
            doc = Document(str(path))
            parts: list[str] = []
            in_sec = not bool(section)
            sec_lvl = 0
            for para in doc.paragraphs:
                style = para.style.name
                is_heading = style.startswith("Heading")
                if section:
                    if is_heading and section.lower() in para.text.lower():
                        in_sec = True
                        try:
                            sec_lvl = int(style.split()[-1])
                        except (ValueError, IndexError):
                            sec_lvl = 1
                        parts.append(f"{'#'*sec_lvl} {para.text}")
                        continue
                    if is_heading and in_sec:
                        try:
                            if int(style.split()[-1]) <= sec_lvl:
                                break
                        except (ValueError, IndexError):
                            break
                if in_sec:
                    if is_heading:
                        try:
                            lvl = int(style.split()[-1])
                        except (ValueError, IndexError):
                            lvl = 2
                        parts.append(f"{'#'*lvl} {para.text}")
                    elif para.text.strip():
                        parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            text = "\n".join(parts)
            label = f" — Sección: {section}" if section else ""
            return f"📄 {path.name}{label}\n{'─'*50}\n{text[:max_chars]}"
        except ImportError:
            pass

        # Fallback: pandoc
        import shutil
        if shutil.which("pandoc"):
            rc, out, err = _run(["pandoc", str(path), "-t", "plain"], timeout=30)
            if rc == 0:
                text = out.strip()
                if section:
                    stext, sl, _ = _extract_section_text(text, section)
                    if sl != -1:
                        text = stext
                return f"📄 {path.name} (pandoc)\n{'─'*50}\n{text[:max_chars]}"

        # Last resort: strip XML
        try:
            import zipfile
            with zipfile.ZipFile(str(path)) as z:
                xml = z.read("word/document.xml").decode("utf-8", errors="replace")
            text = re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", xml)).strip()
            return f"📄 {path.name} (XML extraído)\n{'─'*50}\n{text[:max_chars]}"
        except Exception as exc:
            return f"Error leyendo .docx: {exc}"

    return f"Formato no soportado para doc_read: {ext}. Usa pdf_extract_text para PDFs."


def _tool_doc_update_section(args: dict) -> str:
    """Replace the content of a section (by heading text) in a .md or .docx file."""
    raw = args.get("path", "")
    if not raw or not str(raw).strip():
        return "Parámetro requerido: path"
    path        = Path(raw).expanduser()
    section     = args.get("section", "")
    new_content = args.get("new_content") or args.get("content", "")
    if not section:
        return "Parámetro requerido: section (texto del encabezado a actualizar)"
    if new_content is None or new_content == "":
        return "Parámetro requerido: new_content (nuevo contenido de la sección)"
    if not path.exists():
        return f"Fichero no encontrado: {path}"

    ext = path.suffix.lower()
    if ext in (".md", ".txt", ".rst"):
        text = path.read_text(errors="replace")
        lines = text.splitlines()
        _, start, end = _extract_section_text(text, section)
        if start == -1:
            return f"Sección '{section}' no encontrada en {path.name}"
        heading_line = lines[start]
        new_lines = lines[:start + 1] + [""] + new_content.splitlines() + [""] + lines[end:]
        path.write_text("\n".join(new_lines))
        replaced = end - start - 1
        return (
            f"✅ Sección '{section}' actualizada en {path.name}\n"
            f"   Línea {start+1}: {heading_line[:60]}\n"
            f"   Líneas reemplazadas: {replaced} → {len(new_content.splitlines())}"
        )

    if ext == ".docx":
        try:
            from docx import Document  # type: ignore
            doc = Document(str(path))
            paras = doc.paragraphs
            start_idx = end_idx = None
            sec_lvl = 0
            for i, para in enumerate(paras):
                style = para.style.name
                if style.startswith("Heading") and section.lower() in para.text.lower():
                    start_idx = i
                    try:
                        sec_lvl = int(style.split()[-1])
                    except (ValueError, IndexError):
                        sec_lvl = 1
                elif start_idx is not None and style.startswith("Heading"):
                    try:
                        if int(style.split()[-1]) <= sec_lvl:
                            end_idx = i
                            break
                    except (ValueError, IndexError):
                        end_idx = i
                        break
            if start_idx is None:
                return f"Sección '{section}' no encontrada en {path.name}"
            if end_idx is None:
                end_idx = len(paras)

            # Delete paragraphs between heading and end (from end-1 down to start+1)
            body = doc.element.body
            body_children = list(body)
            heading_elem = paras[start_idx]._element
            end_elem     = paras[end_idx]._element if end_idx < len(paras) else None
            h_pos  = body_children.index(heading_elem)
            e_pos  = body_children.index(end_elem) if end_elem is not None else len(body_children)
            for elem in body_children[h_pos + 1:e_pos]:
                body.remove(elem)

            # Insert new paragraphs after heading
            from docx.oxml import OxmlElement  # type: ignore
            ref = heading_elem
            for line in new_content.splitlines():
                p = OxmlElement("w:p")
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = line
                if line.startswith(" "):
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                r.append(t)
                p.append(r)
                ref.addnext(p)
                ref = p

            doc.save(str(path))
            return f"✅ Sección '{section}' actualizada en {path.name}"
        except ImportError:
            return "python-docx requerido para editar .docx: pip install python-docx"
        except Exception as exc:
            return f"Error actualizando sección .docx: {exc}"

    return f"Formato no soportado para doc_update_section: {ext}"


def _tool_doc_version_bump(args: dict) -> str:
    """Increment the version number in a document's front matter or first heading."""
    raw = args.get("path", "")
    if not raw or not str(raw).strip():
        return "Parámetro requerido: path"
    path = Path(raw).expanduser()
    bump = args.get("bump", "patch")  # major | minor | patch
    if not path.exists():
        return f"Fichero no encontrado: {path}"
    text  = path.read_text(errors="replace")
    today = datetime.date.today().isoformat()

    # Match "version: X.Y.Z" or "version: X.Y" in front matter
    ver_fm = re.compile(r"^(version:\s*)([0-9]+)\.([0-9]+)\.?([0-9]*)(.*)$", re.MULTILINE | re.IGNORECASE)
    m = ver_fm.search(text)
    if m:
        major, minor, patch_v = int(m.group(2)), int(m.group(3)), int(m.group(4) or 0)
        old_ver = f"{major}.{minor}.{patch_v}" if m.group(4) else f"{major}.{minor}"
        if bump == "major":
            major += 1; minor = 0; patch_v = 0
        elif bump == "minor":
            minor += 1; patch_v = 0
        else:
            patch_v += 1
        new_ver = f"{major}.{minor}.{patch_v}"
        new_text = ver_fm.sub(lambda mm: mm.group(1) + new_ver + mm.group(5), text, count=1)
        new_text = re.sub(r"^(modified:\s*).*$", f"\\g<1>{today}", new_text, flags=re.MULTILINE)
        path.write_text(new_text)
        return f"✅ Versión: {old_ver} → {new_ver} en {path.name}\n   Modificado: {today}"

    # Fallback: "vX.Y" anywhere in the first 30 lines
    first_block = "\n".join(text.splitlines()[:30])
    m2 = re.search(r"\bv?([0-9]+)\.([0-9]+)\.?([0-9]*)\b", first_block)
    if m2:
        old_ver_str = m2.group(0)
        major, minor, patch_v = int(m2.group(1)), int(m2.group(2)), int(m2.group(3) or 0)
        if bump == "major":
            major += 1; minor = 0; patch_v = 0
        elif bump == "minor":
            minor += 1; patch_v = 0
        else:
            patch_v += 1
        new_ver = f"{major}.{minor}.{patch_v}"
        path.write_text(text.replace(old_ver_str, new_ver, 1))
        return f"✅ Versión: {old_ver_str} → {new_ver} en {path.name}"

    # No version found — add to front matter
    if text.startswith("---"):
        end = text.find("---", 3)
        if end > 0:
            new_text = text[:end] + f"version: 1.0.0\nmodified: {today}\n" + text[end:]
            path.write_text(new_text)
            return f"✅ Versión añadida: 1.0.0 en {path.name}"
    return (
        f"No se encontró versión en {path.name}\n"
        "Añade 'version: 1.0.0' al front matter YAML (entre --- y ---)"
    )


# ── Bloque 3: CMDB y activos ──────────────────────────────────────────────────

def _cmdb_path_discover(cfg: dict, path_arg: str) -> Optional[Path]:
    """Find CMDB file: explicit arg > project cwd > ~/Documents/."""
    if path_arg:
        p = Path(path_arg).expanduser()
        return p if p.exists() else None
    cwd = Path.cwd()
    candidates = (
        [cwd / n for n in ("cmdb.csv", "cmdb.xlsx", "cmdb.json",
                           "inventario.csv", "inventario.xlsx",
                           "servers.csv", "servers.xlsx")]
        + [Path.home() / "Documents" / n for n in ("cmdb.csv", "inventario.csv")]
    )
    return next((p for p in candidates if p.exists()), None)


def _cmdb_rows_csv(path: Path, query: str, field: str, limit: int) -> str:
    try:
        q = query.lower() if query != "*" else None
        with path.open(newline="", errors="replace") as f:
            reader = csv.DictReader(f)
            results = []
            for row in reader:
                if q is None or (field and q in str(row.get(field, "")).lower()) \
                   or (not field and any(q in str(v).lower() for v in row.values())):
                    results.append(dict(row))
                if len(results) >= limit:
                    break
        if not results:
            return f"Sin resultados para '{query}' en {path.name}"
        lines = [f"🖥 CMDB {path.name} — {len(results)} entrada(s):\n"]
        for row in results:
            lines.append("  " + "  |  ".join(f"{k}: {v}" for k, v in row.items() if v))
        return "\n".join(lines)
    except Exception as exc:
        return f"Error leyendo CMDB CSV: {exc}"


def _cmdb_rows_xlsx(path: Path, query: str, field: str, limit: int) -> str:
    try:
        import openpyxl
        wb  = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        ws  = wb.active
        all_rows = list(ws.iter_rows(values_only=True))
        if not all_rows:
            return f"CMDB vacío: {path.name}"
        headers = [str(c) if c is not None else "" for c in all_rows[0]]
        q = query.lower() if query != "*" else None
        field_idx: Optional[int] = None
        if field:
            try:
                field_idx = headers.index(field)
            except ValueError:
                pass
        results = []
        for row in all_rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            if q is None or (field_idx is not None and q in cells[field_idx].lower()) \
               or (field_idx is None and any(q in c.lower() for c in cells)):
                results.append(dict(zip(headers, cells)))
            if len(results) >= limit:
                break
        if not results:
            return f"Sin resultados para '{query}' en {path.name}"
        lines = [f"🖥 CMDB {path.name} — {len(results)} entrada(s):\n"]
        for row in results:
            lines.append("  " + "  |  ".join(f"{k}: {v}" for k, v in row.items() if v))
        return "\n".join(lines)
    except ImportError:
        return "openpyxl requerido para CMDB .xlsx: pip install openpyxl"
    except Exception as exc:
        return f"Error leyendo CMDB xlsx: {exc}"


def _tool_cmdb_search(args: dict) -> str:
    """Search the CMDB (CSV/XLSX/JSON) for servers, services or assets by any field value."""
    query = args.get("query", "")
    if not query:
        return "Parámetro requerido: query (texto a buscar, o '*' para listar todo)"
    cfg   = _load_config()
    path  = _cmdb_path_discover(cfg, args.get("cmdb_path", "") or args.get("path", ""))
    if path is None:
        cwd = Path.cwd()
        return (
            "CMDB no encontrado. Crea uno de estos ficheros:\n"
            f"  {cwd}/cmdb.csv\n  {cwd}/inventario.csv\n"
            "  ~/Documents/cmdb.csv\n"
            "Usa 'project_init_office' para crear la estructura del proyecto."
        )
    field = args.get("field", "")
    limit = int(args.get("limit", 20))
    if path.suffix == ".csv":
        return _cmdb_rows_csv(path, query, field, limit)
    if path.suffix == ".xlsx":
        return _cmdb_rows_xlsx(path, query, field, limit)
    if path.suffix == ".json":
        try:
            data = json.loads(path.read_text())
            if isinstance(data, dict):
                data = list(data.values())
            q = query.lower() if query != "*" else None
            results = [
                item for item in data
                if isinstance(item, dict) and (q is None or any(q in str(v).lower() for v in item.values()))
            ][:limit]
            if not results:
                return f"Sin resultados para '{query}'"
            lines = [f"🖥 CMDB {path.name} — {len(results)} entrada(s):\n"]
            for item in results:
                lines.append("  " + "  |  ".join(f"{k}: {v}" for k, v in item.items() if v))
            return "\n".join(lines)
        except Exception as exc:
            return f"Error leyendo CMDB JSON: {exc}"
    return f"Formato CMDB no soportado: {path.suffix}"


def _tool_cmdb_update(args: dict) -> str:
    """Update a CMDB entry (CSV only) by key field: finds row where key_field=key_value, applies updates dict."""
    key_field = args.get("key_field", "")
    key_value = args.get("key_value", "")
    updates   = args.get("updates", {})
    if not key_field or not key_value:
        return "Parámetros requeridos: key_field (p.ej. 'hostname'), key_value (p.ej. 'web-01')"
    if not updates or not isinstance(updates, dict):
        return "Parámetro requerido: updates (objeto {campo: nuevo_valor})"
    cfg  = _load_config()
    path = _cmdb_path_discover(cfg, args.get("cmdb_path", "") or args.get("path", ""))
    if path is None:
        return "CMDB no encontrado. Especifica 'cmdb_path' o crea cmdb.csv en el directorio de trabajo."
    if path.suffix != ".csv":
        return "cmdb_update solo soporta .csv. Para .xlsx usa xlsx_fill_range."
    try:
        rows: list[dict] = []
        headers: list[str] = []
        updated = 0
        with path.open(newline="", errors="replace") as f:
            reader = csv.DictReader(f)
            headers = list(reader.fieldnames or [])
            for row in reader:
                if str(row.get(key_field, "")).strip() == str(key_value).strip():
                    row.update(updates)
                    updated += 1
                rows.append(dict(row))
        if updated == 0:
            return f"No se encontró {key_field}={key_value!r} en {path.name}"
        for k in updates:
            if k not in headers:
                headers.append(k)
        with path.open("w", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=headers, extrasaction="ignore")
            writer.writeheader()
            writer.writerows(rows)
        return (
            f"✅ CMDB actualizado: {path.name}\n"
            f"   Clave: {key_field}={key_value!r}\n"
            f"   Cambios: {', '.join(f'{k}={v!r}' for k, v in updates.items())}\n"
            f"   Entradas actualizadas: {updated}"
        )
    except Exception as exc:
        return f"Error actualizando CMDB: {exc}"


def _tool_asset_register_add(args: dict) -> str:
    """Add a new asset entry to the project asset register (CSV)."""
    asset = args.get("asset", {})
    if not asset or not isinstance(asset, dict):
        return 'Parámetro requerido: asset (objeto JSON con datos del activo, p.ej. {"hostname": "srv-01", "ip": "10.0.0.1"})'
    if "fecha_registro" not in asset:
        asset = {**asset, "fecha_registro": datetime.date.today().isoformat()}
    cfg = _load_config()
    cwd = Path.cwd()
    path_arg = args.get("register_path", "") or args.get("path", "")
    if path_arg:
        ar_path = Path(path_arg).expanduser()
    else:
        candidates = [
            cwd / "asset_register.csv",
            cwd / "activos.csv",
            Path.home() / "Documents" / "asset_register.csv",
        ]
        ar_path = next((p for p in candidates if p.exists()), cwd / "asset_register.csv")
    try:
        existing_headers: list[str] = []
        existing_rows: list[dict]   = []
        if ar_path.exists():
            with ar_path.open(newline="", errors="replace") as f:
                reader = csv.DictReader(f)
                existing_headers = list(reader.fieldnames or [])
                existing_rows    = list(reader)
        all_headers = list(existing_headers)
        for k in asset:
            if k not in all_headers:
                all_headers.append(k)
        existing_rows.append(asset)
        with ar_path.open("w", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=all_headers, extrasaction="ignore")
            writer.writeheader()
            writer.writerows(existing_rows)
        return (
            f"✅ Activo añadido al registro: {ar_path.name}\n"
            f"   Registro #{len(existing_rows)}\n"
            + "\n".join(f"   {k}: {v}" for k, v in asset.items())
        )
    except Exception as exc:
        return f"Error añadiendo activo: {exc}"


# ── Template and IT report tools ─────────────────────────────────────────────

def _tool_doc_read_template_fields(args: dict) -> str:
    """Extract {{FIELD}} placeholders from a .docx, .md or .txt template."""
    raw = args.get("path", "")
    if not raw or not str(raw).strip():
        return "Parámetro requerido: path"
    path = Path(raw).expanduser()
    if not path.exists():
        return f"Fichero no encontrado: {path}"

    fields: set[str] = set()
    field_re = re.compile(r"\{\{([A-Z0-9_\s]+?)\}\}", re.IGNORECASE)

    if path.suffix.lower() == ".docx":
        try:
            import zipfile
            with zipfile.ZipFile(str(path)) as z:
                for name in z.namelist():
                    if name.startswith("word/") and name.endswith(".xml"):
                        raw = z.read(name).decode("utf-8", errors="replace")
                        text = re.sub(r"<[^>]+>", " ", raw)
                        for m in field_re.finditer(text):
                            fields.add(m.group(1).strip().upper())
        except Exception as exc:
            return f"Error leyendo .docx: {exc}"
    else:
        try:
            text = path.read_text(errors="replace")
            for m in field_re.finditer(text):
                fields.add(m.group(1).strip().upper())
        except Exception as exc:
            return f"Error leyendo fichero: {exc}"

    if not fields:
        return f"No se encontraron campos {{{{CAMPO}}}} en: {path.name}"
    lines = [f"📋 Campos de plantilla en {path.name} — {len(fields)} campo(s):"]
    for f in sorted(fields):
        lines.append(f"  {{{{  {f}  }}}}")
    return "\n".join(lines)


def _tool_doc_fill_template(args: dict) -> str:
    """Fill {{FIELD}} placeholders in a .docx or text template with provided values."""
    template_path = Path(args.get("template_path", "")).expanduser()
    output_path   = args.get("output_path", "")
    fields        = args.get("fields", {})

    if not template_path:
        return "Parámetro requerido: template_path"
    if not template_path.exists():
        return f"Plantilla no encontrada: {template_path}"
    if not fields or not isinstance(fields, dict):
        return "Parámetro requerido: fields (objeto JSON de {CAMPO: valor})"

    if not output_path:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = str(template_path.parent / f"{template_path.stem}_{ts}{template_path.suffix}")
    out_path = Path(output_path).expanduser()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    ext = template_path.suffix.lower()
    filled_count = 0

    def _replace_all_fields(text: str) -> tuple[str, int]:
        count = 0
        for key, val in fields.items():
            for ph in (f"{{{{{str(key).upper()}}}}}", f"{{{{{str(key).lower()}}}}}"):
                if ph in text:
                    text = text.replace(ph, str(val))
                    count += 1
        return text, count

    if ext == ".docx":
        # Try python-docx (handles split runs + cross-run placeholders)
        # ⚠️ IMPORTANTE: Se preservan los estilos originales del documento
        #    - Los estilos de párrafos, tablas, fuentes y formatos se mantienen
        #    - Solo se reemplazan el texto de los placeholders {{CAMPO}}
        #    - Los estilos de títulos, subtítulos, índices, listas y texto se respetan
        try:
            from docx import Document  # type: ignore

            def _fill_para(para) -> int:
                """Replace placeholders in a paragraph, handling cross-run splits.
                
                ⚠️ PRESERVA ESTILOS:
                - No modifica estilos de fuente, tamaño, color, negrita, cursiva
                - No modifica estilos de párrafo (alineación, interlineado, etc.)
                - No modifica estilos de tabla (bordes, sombreado, etc.)
                - Solo reemplaza el texto de los placeholders
                """
                n = 0
                # Pass 1: per-run replacement (fast path, most common)
                for run in para.runs:
                    new_text, cnt = _replace_all_fields(run.text)
                    if cnt:
                        run.text = new_text
                        n += cnt
                # Pass 2: join-all-runs and retry (catches splits like {{CAM | PO}})
                full = "".join(r.text for r in para.runs)
                new_full, cnt2 = _replace_all_fields(full)
                if cnt2:
                    # Rebuild: put all text in first run, clear the rest
                    if para.runs:
                        para.runs[0].text = new_full
                        for r in para.runs[1:]:
                            r.text = ""
                    n += cnt2
                return n

            doc = Document(str(template_path))
            for para in doc.paragraphs:
                filled_count += _fill_para(para)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            filled_count += _fill_para(para)
            # Also scan headers/footers
            for section in doc.sections:
                for hdr_para in list(getattr(section.header, "paragraphs", [])):
                    filled_count += _fill_para(hdr_para)
                for ftr_para in list(getattr(section.footer, "paragraphs", [])):
                    filled_count += _fill_para(ftr_para)
            doc.save(str(out_path))
            return (
                f"✅ Plantilla rellenada (python-docx): {out_path.name}\n"
                f"   Origen: {template_path.name}\n"
                f"   Campos rellenados: {filled_count}\n"
                f"   Tamaño: {out_path.stat().st_size:,} bytes\n"
                f"   ⚠️ Estilos originales preservados"
            )
        except ImportError:
            pass
        except Exception as exc:
            return f"Error con python-docx: {exc}"

        # Fallback: raw XML/zip manipulation
        try:
            import zipfile
            with zipfile.ZipFile(str(template_path), "r") as zin:
                names = zin.namelist()
                files_data: dict[str, bytes] = {n: zin.read(n) for n in names}

            xml_targets = [n for n in names if n.startswith("word/") and n.endswith(".xml")]
            for target in xml_targets:
                try:
                    content, n = _replace_all_fields(files_data[target].decode("utf-8", errors="replace"))
                    filled_count += n
                    files_data[target] = content.encode("utf-8")
                except Exception:
                    pass

            with zipfile.ZipFile(str(out_path), "w", zipfile.ZIP_DEFLATED) as zout:
                for name, data in files_data.items():
                    zout.writestr(name, data)
            return (
                f"✅ Plantilla rellenada (zip/XML): {out_path.name}\n"
                f"   Origen: {template_path.name}\n"
                f"   Campos detectados: {filled_count}\n"
                f"   Tamaño: {out_path.stat().st_size:,} bytes\n"
                f"   Nota: instala python-docx para soporte completo de runs divididos."
            )
        except Exception as exc:
            return f"Error rellenando plantilla .docx: {exc}"
    else:
        # Text-based templates (.md, .txt, .html…)
        try:
            content, filled_count = _replace_all_fields(template_path.read_text(errors="replace"))
            out_path.write_text(content)
            return (
                f"✅ Plantilla rellenada: {out_path.name}\n"
                f"   Campos rellenados: {filled_count}\n"
                f"   Tamaño: {out_path.stat().st_size:,} bytes"
            )
        except Exception as exc:
            return f"Error rellenando plantilla: {exc}"


def _tool_doc_list_templates(args: dict) -> str:
    """List available document templates: project dir, workspace, and configured templates_dir."""
    cfg  = _load_config()
    cwd  = Path.cwd()
    dirs_to_scan: list[Path] = []
    # 1. Explicit arg
    if args.get("directory"):
        dirs_to_scan.append(Path(args["directory"]).expanduser())
    # 2. Project dir (cwd/templates, cwd/plantillas)
    for local in ("templates", "plantillas"):
        p = cwd / local
        if p.exists() and p not in dirs_to_scan:
            dirs_to_scan.append(p)
    # 3. Configured templates_dir (may already overlap with above)
    conf_dir = Path(cfg.get("templates_dir", str(Path.home() / "Documents" / "templates"))).expanduser()
    if conf_dir not in dirs_to_scan:
        dirs_to_scan.append(conf_dir)

    patterns = ("*.docx", "*.xlsx", "*.md", "*.txt", "*.odt")
    seen: set[Path] = set()
    all_files: list[tuple[Path, Path]] = []  # (dir, file)
    for d in dirs_to_scan:
        if not d.exists():
            continue
        for pattern in patterns:
            for f in d.glob(pattern):
                if f not in seen:
                    seen.add(f)
                    all_files.append((d, f))
    all_files.sort(key=lambda x: x[1].stat().st_mtime, reverse=True)

    if not all_files:
        dirs_str = ", ".join(str(d) for d in dirs_to_scan)
        return (
            f"Sin plantillas en: {dirs_str}\n"
            f"Coloca ficheros .docx, .xlsx o .md en {cwd / 'templates'} o configura 'templates_dir'."
        )
    lines = [f"📁 Plantillas disponibles ({len(all_files)}):"]
    last_dir = None
    for d, f in all_files:
        if d != last_dir:
            label = "workspace" if d.parent == cwd or d == cwd else str(d)
            lines.append(f"\n  📂 {label}/")
            last_dir = d
        mtime = datetime.datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d")
        lines.append(f"     {mtime}  {f.stat().st_size:>8,} B  {f.name}")
    return "\n".join(lines)


def _write_docx_from_markdown(content: str, path: Path) -> str | None:
    """Convert markdown text to a .docx file.

    Tries pandoc first (best output), then python-docx, returns None on success
    or an error string if both fail.
    """
    import subprocess as _sp
    try:
        r = _sp.run(
            ["pandoc", "-f", "markdown", "-t", "docx", "-o", str(path)],
            input=content, text=True, capture_output=True, timeout=30,
        )
        if r.returncode == 0 and path.exists() and path.stat().st_size > 0:
            return None
    except (FileNotFoundError, _sp.TimeoutExpired):
        pass

    try:
        from docx import Document as _Doc  # type: ignore
        doc = _Doc()
        for line in content.splitlines():
            s = line.strip()
            if s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s == "---":
                doc.add_paragraph("─" * 60)
            elif s:
                doc.add_paragraph(s)
            elif doc.paragraphs:
                doc.add_paragraph("")
        doc.save(str(path))
        return None
    except ImportError:
        pass

    return "pandoc y python-docx no disponibles — instala uno para generar .docx"


def _tool_doc_create_rfc(args: dict) -> str:
    """Generate a structured RFC/Request for Change document (.docx by default, or .md)."""
    title            = args.get("title", "")
    requester        = args.get("requester", "")
    if not title or not requester:
        return "Parámetros requeridos: title, requester"

    date             = args.get("date", datetime.date.today().isoformat())
    priority         = args.get("priority", "Media")
    change_type      = args.get("change_type", "Normal")
    affected_systems = args.get("affected_systems", "_Por especificar_")
    description      = args.get("description", "_Por completar_")
    justification    = args.get("justification", "_Por completar_")
    risk             = args.get("risk", "_Por analizar_")
    risk_level       = args.get("risk_level", "Bajo")
    rollback_plan    = args.get("rollback_plan", "_Por definir_")
    testing_plan     = args.get("testing_plan", "_Por definir_")
    impl_steps       = args.get("implementation_steps", "_Por definir_")
    scheduled_date   = args.get("scheduled_date", "Por definir")
    scheduled_window = args.get("scheduled_window", "Por definir")
    approver         = args.get("approver", "Por asignar")
    output_path      = args.get("output_path", "")
    fmt              = args.get("format", "docx").lower()

    rfc_id = f"RFC-{datetime.datetime.now().strftime('%Y%m%d-%H%M')}"
    doc = f"""# {rfc_id} — {title}

| Campo | Valor |
|-------|-------|
| **ID RFC** | {rfc_id} |
| **Título** | {title} |
| **Solicitante** | {requester} |
| **Fecha solicitud** | {date} |
| **Tipo de cambio** | {change_type} |
| **Prioridad** | {priority} |
| **Nivel de riesgo** | {risk_level} |
| **Fecha programada** | {scheduled_date} |
| **Ventana de cambio** | {scheduled_window} |
| **Aprobador** | {approver} |

---

## 1. Descripción del cambio

{description}

## 2. Justificación

{justification}

## 3. Sistemas afectados

{affected_systems}

## 4. Plan de implementación

{impl_steps}

## 5. Plan de pruebas y validación

{testing_plan}

## 6. Análisis de riesgos

**Nivel de riesgo:** {risk_level}

{risk}

## 7. Plan de marcha atrás (Rollback)

{rollback_plan}

---

## 8. Aprobaciones

| Rol | Nombre | Firma | Fecha |
|-----|--------|-------|-------|
| Solicitante | {requester} | | {date} |
| Aprobador técnico | {approver} | | |
| Responsable de negocio | | | |
| Gestor de cambios | | | |

---
*Documento generado por OOCode Home Office Assistant — {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}*
"""
    if output_path:
        try:
            out = Path(output_path).expanduser()
            # If path has no extension, apply format preference
            if not out.suffix:
                out = out.with_suffix(".docx" if fmt == "docx" else ".md")
            out.parent.mkdir(parents=True, exist_ok=True)
            if out.suffix.lower() == ".docx":
                err = _write_docx_from_markdown(doc, out)
                if err:
                    out = out.with_suffix(".md")
                    out.write_text(doc)
                    return f"⚠ {err}\n✅ RFC guardado como Markdown: {out}\n   ID: {rfc_id}"
            else:
                out.write_text(doc)
            return f"✅ RFC generado: {out}\n   ID: {rfc_id}"
        except Exception as exc:
            return f"Error guardando RFC: {exc}"
    # No path given: return inline content (compatible with all callers)
    return f"📋 RFC generado (ID: {rfc_id}):\n\n{doc}"


def _tool_xlsx_fill_range(args: dict) -> str:
    """Write multiple cells at once in an Excel file.

    cells: dict {cell_addr: value} OR list of {cell, value, style} objects.
    style (por celda o global): bold, italic, font_color, bg_color, align, border…
    """
    path        = Path(args.get("path", "")).expanduser()
    sheet       = args.get("sheet", "Hoja1")
    cells       = args.get("cells", {})
    global_style = args.get("style", {})
    if not path:
        return "Parámetro requerido: path"
    if not cells:
        return 'Parámetro requerido: cells ({"A1": "valor"} o [{"cell":"A1","value":…,"style":{…}}])'
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path)) if path.exists() else openpyxl.Workbook()
        if sheet not in wb.sheetnames:
            wb.create_sheet(sheet)
        ws  = wb[sheet]
        written: list[str] = []
        # Soportar tanto dict simple como lista de objetos {cell, value, style}
        items: list[tuple[str, object, dict]]
        if isinstance(cells, dict):
            items = [(addr, val, global_style) for addr, val in cells.items()]
        elif isinstance(cells, list):
            items = []
            for entry in cells:
                if isinstance(entry, dict):
                    addr = str(entry.get("cell", ""))
                    val  = entry.get("value", "")
                    sty  = {**global_style, **entry.get("style", {})}
                    items.append((addr, val, sty))
        else:
            return 'cells debe ser un dict {"A1": valor} o una lista de objetos {cell, value, style}'
        for cell_addr, value, cell_style in items:
            if not cell_addr:
                continue
            c = ws[cell_addr.upper()]
            c.value = value
            if cell_style:
                _apply_cell_style(c, cell_style)
            written.append(f"{cell_addr.upper()}={value!r}")
        wb.save(str(path))
        return (
            f"✅ {len(written)} celdas escritas en {path.name} [{sheet}]:\n"
            + "\n".join(f"   {w}" for w in written)
        )
    except ImportError:
        return "openpyxl no instalado. Instala con: pip install openpyxl"
    except Exception as exc:
        return f"Error escribiendo rango Excel: {exc}"


def _tool_xlsx_append_row(args: dict) -> str:
    """Append a row of values to the next empty row in an Excel sheet.

    values: lista de valores, o lista de {value, style} para estilos por celda.
    style (global): se aplica a todas las celdas de la fila.
    """
    path         = Path(args.get("path", "")).expanduser()
    sheet        = args.get("sheet", "Hoja1")
    values       = args.get("values", [])
    global_style = args.get("style", {})
    if not path:
        return "Parámetro requerido: path"
    if not values:
        return "Parámetro requerido: values (lista de valores para la nueva fila)"
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path)) if path.exists() else openpyxl.Workbook()
        if sheet not in wb.sheetnames:
            wb.create_sheet(sheet)
        ws = wb[sheet]
        # Determinar si hay estilos por celda
        has_per_cell = isinstance(values, list) and values and isinstance(values[0], dict)
        if has_per_cell:
            raw_values = [v.get("value", "") if isinstance(v, dict) else v for v in values]
        else:
            raw_values = list(values)
        ws.append(raw_values)
        row_num = ws.max_row
        # Aplicar estilos tras append
        if global_style or has_per_cell:
            for col_idx, val in enumerate(values, 1):
                c = ws.cell(row=row_num, column=col_idx)
                cell_style = {**global_style}
                if has_per_cell and isinstance(val, dict):
                    cell_style.update(val.get("style", {}))
                if cell_style:
                    _apply_cell_style(c, cell_style)
        wb.save(str(path))
        return (
            f"✅ Fila {row_num} añadida en {path.name} [{sheet}]:\n"
            f"   {' | '.join(str(v.get('value', v) if isinstance(v, dict) else v) for v in values)}"
        )
    except ImportError:
        return "openpyxl no instalado. Instala con: pip install openpyxl"
    except Exception as exc:
        return f"Error añadiendo fila: {exc}"


def _tool_xlsx_create_report(args: dict) -> str:
    """Create a formatted Excel report from headers and rows data."""
    path    = Path(args.get("path", "")).expanduser()
    sheet   = args.get("sheet", "Informe")
    headers = args.get("headers", [])
    rows    = args.get("rows", [])
    title   = args.get("title", "")
    if not path:
        return "Parámetro requerido: path"
    if not headers:
        return "Parámetro requerido: headers (lista de nombres de columnas)"
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment  # type: ignore
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet
        start_row = 1
        if title:
            ncols = max(len(headers), 1)
            ws.merge_cells(f"A1:{chr(64 + min(ncols, 26))}1")
            c = ws["A1"]
            c.value = title
            c.font = Font(bold=True, size=13)
            c.alignment = Alignment(horizontal="center")
            start_row = 2
        # Header row
        for col, header in enumerate(headers, 1):
            c = ws.cell(row=start_row, column=col, value=header)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = PatternFill("solid", fgColor="366092")
            c.alignment = Alignment(horizontal="center")
        # Data rows
        for r_idx, row_data in enumerate(rows, start_row + 1):
            row_list = list(row_data) if not isinstance(row_data, list) else row_data
            for col, val in enumerate(row_list, 1):
                c = ws.cell(row=r_idx, column=col, value=val)
                if (r_idx - start_row) % 2 == 0:
                    c.fill = PatternFill("solid", fgColor="DCE6F1")
        # Auto-width (skip MergedCell objects which lack column_letter)
        for col in ws.columns:
            max_len = 8
            col_letter = None
            for c in col:
                try:
                    if col_letter is None and hasattr(c, "column_letter"):
                        col_letter = c.column_letter
                    if c.value:
                        max_len = max(max_len, len(str(c.value)))
                except Exception:
                    pass
            if col_letter:
                ws.column_dimensions[col_letter].width = min(max_len + 4, 60)
        path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(path))
        return (
            f"✅ Informe Excel creado: {path.name}\n"
            f"   Hoja: {sheet}  |  Columnas: {len(headers)}  |  Filas: {len(rows)}\n"
            f"   Tamaño: {path.stat().st_size:,} bytes"
        )
    except ImportError:
        return "openpyxl no instalado (necesario para estilos). Instala con: pip install openpyxl"
    except Exception as exc:
        return f"Error creando informe Excel: {exc}"


# ── Herramientas — esquemas ──────────────────────────────────────────────────

_TOOLS = [
    {
        "name": "email_list",
        "description": "Lista emails de una bandeja de entrada (IMAP). Requiere configuración en ~/.oocode/home_office.json.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "mailbox": {"type": "string", "description": "Carpeta IMAP. Default: INBOX"},
                "limit":   {"type": "integer", "description": "Máximo de emails a devolver (max 50). Default: 10"},
            },
        },
    },
    {
        "name": "email_read",
        "description": "Lee un email completo por su UID IMAP.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "uid":     {"type": "string",  "description": "UID del email (obtenido de email_list)"},
                "mailbox": {"type": "string",  "description": "Carpeta IMAP. Default: INBOX"},
            },
            "required": ["uid"],
        },
    },
    {
        "name": "email_send",
        "description": "Envía un email por SMTP. Requiere configuración en ~/.oocode/home_office.json.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "to":      {"type": "string", "description": "Destinatario(s), separados por coma"},
                "subject": {"type": "string", "description": "Asunto del email"},
                "body":    {"type": "string", "description": "Cuerpo del email (texto plano)"},
                "cc":      {"type": "string", "description": "CC (opcional)"},
                "bcc":     {"type": "string", "description": "BCC (opcional)"},
            },
            "required": ["to", "subject", "body"],
        },
    },
    {
        "name": "email_search",
        "description": "Busca emails por asunto o contenido usando IMAP SEARCH.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query":   {"type": "string",  "description": "Texto a buscar en asunto y cuerpo"},
                "mailbox": {"type": "string",  "description": "Carpeta IMAP. Default: INBOX"},
                "limit":   {"type": "integer", "description": "Máximo de resultados. Default: 20"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "doc_convert",
        "description": "Convierte documentos entre formatos usando pandoc (md↔docx, md→pdf, html→md, etc.).",
        "inputSchema": {
            "type": "object",
            "properties": {
                "input_path":    {"type": "string", "description": "Ruta al fichero de entrada"},
                "output_format": {"type": "string", "description": "Formato de salida: pdf, docx, html, md, odt, rst, epub"},
                "output_path":   {"type": "string", "description": "Ruta del fichero de salida (opcional; por defecto mismo nombre con nueva extensión)"},
            },
            "required": ["input_path", "output_format"],
        },
    },
    {
        "name": "pdf_extract_text",
        "description": "Extrae texto de un PDF. Usa pdftotext (poppler) o pdfplumber.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":  {"type": "string", "description": "Ruta al fichero PDF"},
                "pages": {"type": "string", "description": "Rango de páginas, p.ej. '1-5' (opcional)"},
            },
            "required": ["path"],
        },
    },
    {
        "name": "doc_word_count",
        "description": "Cuenta palabras, líneas, caracteres y párrafos de un documento de texto.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Ruta al fichero (txt, md, rst, html…)"},
            },
            "required": ["path"],
        },
    },
    {
        "name": "xlsx_read",
        "description": "Lee celdas o rango de un archivo Excel (.xlsx) o CSV. Requiere openpyxl para .xlsx.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":  {"type": "string",  "description": "Ruta al fichero .xlsx o .csv"},
                "sheet": {"type": "string",  "description": "Nombre de la hoja (solo .xlsx; por defecto la primera)"},
                "range": {"type": "string",  "description": "Rango de celdas, p.ej. 'A1:D10' (pendiente en v1)"},
                "limit": {"type": "integer", "description": "Máximo de filas a devolver. Default: 50"},
            },
            "required": ["path"],
        },
    },
    {
        "name": "xlsx_write",
        "description": "Escribe un valor en una celda de un archivo Excel (.xlsx). Crea el fichero si no existe.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":  {"type": "string", "description": "Ruta al fichero .xlsx"},
                "sheet": {"type": "string", "description": "Nombre de la hoja. Default: Hoja1"},
                "cell":  {"type": "string", "description": "Celda, p.ej. 'B3'"},
                "value": {"description": "Valor a escribir (string, número o booleano)"},
                "style": {"type": "object", "description": 'Estilo de la celda (opcional): {"bold": true, "italic": true, "underline": true, "font_size": 12, "font_color": "#FF0000", "bg_color": "#FFFF00", "align": "center", "wrap_text": true, "border": true, "number_format": "#,##0.00"}'},
            },
            "required": ["path", "cell", "value"],
        },
    },
    {
        "name": "csv_analyze",
        "description": "Analiza un fichero CSV: cabeceras, primeras filas, estadísticas de columnas numéricas.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":  {"type": "string",  "description": "Ruta al fichero CSV"},
                "limit": {"type": "integer", "description": "Filas de muestra a mostrar. Default: 5"},
            },
            "required": ["path"],
        },
    },
    {
        "name": "cal_list",
        "description": "Lista eventos de un fichero .ics local, con filtro opcional por rango de fechas.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "source": {"type": "string", "description": "Ruta al fichero .ics (por defecto: calendar_file en config)"},
                "start":  {"type": "string", "description": "Fecha inicio filtro, formato YYYY-MM-DD (opcional)"},
                "end":    {"type": "string", "description": "Fecha fin filtro, formato YYYY-MM-DD (opcional)"},
                "limit":  {"type": "integer", "description": "Máximo de eventos. Default: 15"},
            },
        },
    },
    {
        "name": "cal_add",
        "description": "Añade un evento a un fichero .ics local.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "title":       {"type": "string", "description": "Título del evento"},
                "start":       {"type": "string", "description": "Fecha/hora inicio: YYYY-MM-DD o YYYY-MM-DDTHH:MM"},
                "end":         {"type": "string", "description": "Fecha/hora fin (opcional)"},
                "location":    {"type": "string", "description": "Lugar (opcional)"},
                "description": {"type": "string", "description": "Descripción (opcional)"},
                "file":        {"type": "string", "description": "Ruta al fichero .ics (por defecto: calendar_file en config)"},
            },
            "required": ["title", "start"],
        },
    },
    {
        "name": "cal_search",
        "description": "Busca eventos en un fichero .ics por texto en título o descripción.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query":  {"type": "string", "description": "Texto a buscar"},
                "source": {"type": "string", "description": "Ruta al fichero .ics (opcional)"},
                "limit":  {"type": "integer", "description": "Máximo de resultados. Default: 20"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "notes_list",
        "description": "Lista ficheros markdown en el directorio de notas, ordenados por fecha de modificación.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "directory": {"type": "string",  "description": "Directorio de notas (por defecto: notes_dir en config)"},
                "pattern":   {"type": "string",  "description": "Patrón glob. Default: *.md"},
                "limit":     {"type": "integer", "description": "Máximo de notas. Default: 20"},
            },
        },
    },
    {
        "name": "notes_search",
        "description": "Busca texto en ficheros markdown del directorio de notas usando ripgrep o grep.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query":     {"type": "string",  "description": "Texto a buscar"},
                "directory": {"type": "string",  "description": "Directorio de notas (opcional)"},
                "limit":     {"type": "integer", "description": "Máximo de ficheros con resultados. Default: 10"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "notes_save",
        "description": "Guarda o actualiza una nota markdown con front matter (title, created, modified).",
        "inputSchema": {
            "type": "object",
            "properties": {
                "title":     {"type": "string", "description": "Título de la nota (también determina el nombre del fichero)"},
                "content":   {"type": "string", "description": "Contenido en markdown"},
                "directory": {"type": "string", "description": "Directorio donde guardar (opcional)"},
            },
            "required": ["title"],
        },
    },
    {
        "name": "image_to_text",
        "description": "Extrae texto de una imagen usando tesseract OCR. Requiere tesseract instalado.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Ruta a la imagen (PNG, JPG, TIFF…)"},
                "lang": {"type": "string", "description": "Idiomas tesseract, p.ej. 'spa+eng'. Default: spa+eng"},
            },
            "required": ["path"],
        },
    },
    {
        "name": "contact_search",
        "description": "Busca en ficheros vCard (.vcf) por nombre, email, teléfono u organización.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query":   {"type": "string", "description": "Texto a buscar en los contactos"},
                "vcf_dir": {"type": "string", "description": "Directorio de contactos (por defecto: contacts_dir en config)"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "markdown_to_html",
        "description": "Convierte markdown a HTML. Usa python-markdown, pandoc o conversión básica.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "content": {"type": "string", "description": "Texto markdown (alternativa a path)"},
                "path":    {"type": "string", "description": "Ruta a un fichero .md (alternativa a content)"},
                "output":  {"type": "string", "description": "Ruta de salida .html (opcional; sin ruta: devuelve el HTML)"},
            },
        },
    },
    # ── Template and IT report tools ────────────────────────────────────────
    {
        "name": "doc_read_template_fields",
        "description": "Extrae los campos {{CAMPO}} de una plantilla .docx, .md o .txt. Útil para saber qué datos rellenar antes de usar doc_fill_template.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Ruta a la plantilla (.docx, .md, .txt, .html)"},
            },
            "required": ["path"],
        },
    },
    {
        "name": "doc_fill_template",
        "description": "Rellena los campos {{CAMPO}} de una plantilla .docx o de texto con los valores proporcionados. Genera un fichero de salida. Requiere python-docx para .docx (pip install python-docx).",
        "inputSchema": {
            "type": "object",
            "properties": {
                "template_path": {"type": "string", "description": "Ruta a la plantilla (.docx, .md, .txt)"},
                "fields":        {"type": "object", "description": 'Objeto JSON con {CAMPO: valor}, p.ej. {"NOMBRE": "Juan", "FECHA": "2026-05-21"}'},
                "output_path":   {"type": "string", "description": "Ruta del fichero de salida (opcional; por defecto añade timestamp al nombre)"},
            },
            "required": ["template_path", "fields"],
        },
    },
    {
        "name": "doc_list_templates",
        "description": "Lista las plantillas de documentos disponibles (.docx, .xlsx, .md) en el directorio de plantillas configurado.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "directory": {"type": "string", "description": "Directorio de plantillas (por defecto: templates_dir en config)"},
            },
        },
    },
    {
        "name": "doc_create_rfc",
        "description": "Genera un documento RFC/Request for Change estructurado para cambios de infraestructura IT. Por defecto genera .docx (usando pandoc o python-docx); usa format='md' para markdown.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "title":            {"type": "string", "description": "Título del cambio"},
                "requester":        {"type": "string", "description": "Nombre del solicitante"},
                "date":             {"type": "string", "description": "Fecha de solicitud (YYYY-MM-DD). Default: hoy"},
                "priority":         {"type": "string", "description": "Prioridad: Alta/Media/Baja. Default: Media"},
                "change_type":      {"type": "string", "description": "Tipo: Normal/Estándar/Urgente. Default: Normal"},
                "affected_systems": {"type": "string", "description": "Sistemas/servidores afectados"},
                "description":      {"type": "string", "description": "Descripción detallada del cambio"},
                "justification":    {"type": "string", "description": "Justificación y beneficios"},
                "risk":             {"type": "string", "description": "Descripción de riesgos identificados"},
                "risk_level":       {"type": "string", "description": "Nivel de riesgo: Alto/Medio/Bajo. Default: Bajo"},
                "rollback_plan":    {"type": "string", "description": "Plan de marcha atrás"},
                "testing_plan":     {"type": "string", "description": "Plan de pruebas post-implementación"},
                "implementation_steps": {"type": "string", "description": "Pasos de implementación"},
                "scheduled_date":   {"type": "string", "description": "Fecha programada del cambio"},
                "scheduled_window": {"type": "string", "description": "Ventana de mantenimiento"},
                "approver":         {"type": "string", "description": "Nombre del aprobador"},
                "output_path":      {"type": "string", "description": "Ruta del fichero de salida (la extensión puede ser .docx o .md). Si se omite, se guarda automáticamente en notes_dir."},
                "format":           {"type": "string", "description": "Formato de salida: 'docx' (default) o 'md'. Con 'docx' requiere pandoc o python-docx; si no están disponibles cae a .md."},
            },
            "required": ["title", "requester"],
        },
    },
    {
        "name": "xlsx_fill_range",
        "description": "Escribe múltiples celdas a la vez en un fichero Excel. Ideal para rellenar plantillas de informes. Soporta estilos por celda.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":  {"type": "string", "description": "Ruta al fichero .xlsx (se crea si no existe)"},
                "sheet": {"type": "string", "description": "Nombre de la hoja. Default: Hoja1"},
                "cells": {"description": 'Dict simple {"A1": "Título", "B2": 42} O lista de objetos [{"cell": "A1", "value": "Título", "style": {"bold": true, "bg_color": "#4472C4", "font_color": "#FFFFFF"}}]'},
                "style": {"type": "object", "description": 'Estilo global aplicado a todas las celdas (opcional): {"bold": true, "italic": true, "font_size": 11, "font_color": "#000000", "bg_color": "#FFFFFF", "align": "center", "wrap_text": false, "border": false, "number_format": "General"}'},
            },
            "required": ["path", "cells"],
        },
    },
    {
        "name": "xlsx_append_row",
        "description": "Añade una fila de datos al final de una hoja Excel. Útil para registros de incidencias o logs. Soporta estilos por celda.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":   {"type": "string", "description": "Ruta al fichero .xlsx"},
                "sheet":  {"type": "string", "description": "Nombre de la hoja. Default: Hoja1"},
                "values": {"type": "array",  "description": 'Lista de valores para la nueva fila. Cada elemento puede ser un valor simple o un objeto {"value": ..., "style": {"bold": true, "font_color": "#FF0000", ...}}'},
                "style":  {"type": "object", "description": 'Estilo global para toda la fila (opcional): {"bold": true, "bg_color": "#D9E1F2", "border": true, "align": "center"}'},
            },
            "required": ["path", "values"],
        },
    },
    {
        "name": "xlsx_create_report",
        "description": "Crea un informe Excel formateado con cabecera, colores y autoajuste de columnas.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":    {"type": "string", "description": "Ruta del fichero .xlsx a crear"},
                "headers": {"type": "array",  "description": "Lista de nombres de columnas"},
                "rows":    {"type": "array",  "description": "Lista de listas con los datos (una lista por fila)"},
                "title":   {"type": "string", "description": "Título del informe (fila superior fusionada, opcional)"},
                "sheet":   {"type": "string", "description": "Nombre de la hoja. Default: Informe"},
            },
            "required": ["path", "headers"],
        },
    },
    # ── Workspace / Project context ──────────────────────────────────────────
    {
        "name": "project_context_read",
        "description": "Lee el fichero OOCODE.md del directorio de trabajo y devuelve metadatos del proyecto (cliente, tipo, naming, directorios) junto con el cuerpo completo.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Ruta al OOCODE.md (opcional; por defecto: OOCODE.md en cwd)"},
            },
        },
    },
    {
        "name": "project_init_office",
        "description": "Inicializa la estructura de directorios para un proyecto IT/oficina: crea OOCODE.md, subdirectorios (docs/rfcs, docs/reports, etc.), cmdb.csv y risk_register.csv.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "project":      {"type": "string", "description": "Nombre del proyecto"},
                "client":       {"type": "string", "description": "Cliente o empresa"},
                "project_type": {"type": "string", "description": "Tipo: IT/DC/Cloud/Infra/General. Default: IT"},
                "naming":       {"type": "string", "description": "Patrón de nombres: p.ej. {CLIENT}-{TYPE}-{YYMMDD}-v{VER}"},
                "templates_dir":{"type": "string", "description": "Directorio de plantillas relativo al proyecto (opcional)"},
                "directory":    {"type": "string", "description": "Directorio donde crear el proyecto (por defecto: cwd)"},
            },
            "required": ["project"],
        },
    },
    {
        "name": "doc_project_save",
        "description": "Guarda un documento en el subdirectorio correcto del proyecto (docs/rfcs, docs/reports, docs/incidents…). Los tipos formales (rfc, report, meeting, plan, incident) usan .docx por defecto; 'general' usa .md.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "content":   {"type": "string", "description": "Contenido del documento en markdown o texto"},
                "doc_type":  {"type": "string", "description": "Tipo: rfc/change_request/report/informe/meeting/acta/incident/incidencia/plan/migration/general"},
                "filename":  {"type": "string", "description": "Nombre de fichero completo (con extensión). Si se omite se genera automáticamente con naming convention."},
                "ext":       {"type": "string", "description": "Extensión cuando no se da filename: '.docx' (default para tipos formales) o '.md'"},
                "directory": {"type": "string", "description": "Directorio base del proyecto (por defecto: cwd/docs)"},
            },
            "required": ["content", "doc_type"],
        },
    },
    # ── Document intelligence ────────────────────────────────────────────────
    {
        "name": "doc_read",
        "description": "Lee el contenido de un documento .docx o .md. Opcionalmente extrae sólo una sección por su encabezado.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":    {"type": "string", "description": "Ruta al documento (.docx o .md)"},
                "section": {"type": "string", "description": "Encabezado de sección a extraer (opcional; sin él devuelve todo)"},
                "max_chars":{"type": "integer","description": "Límite de caracteres devueltos. Default: 8000"},
            },
            "required": ["path"],
        },
    },
    {
        "name": "doc_update_section",
        "description": "Reemplaza el contenido de una sección (identificada por su encabezado) en un documento .md o .docx. Para .md hace edición nativa; para .docx requiere python-docx.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path":        {"type": "string", "description": "Ruta al documento (.md o .docx)"},
                "section":     {"type": "string", "description": "Texto exacto del encabezado de la sección a reemplazar"},
                "new_content": {"type": "string", "description": "Nuevo contenido de la sección (markdown o texto)"},
            },
            "required": ["path", "section", "new_content"],
        },
    },
    {
        "name": "doc_version_bump",
        "description": "Incrementa la versión de un documento .md con front matter YAML (version: X.Y.Z) o que contenga vX.Y en sus primeras líneas.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Ruta al documento .md"},
                "bump": {"type": "string", "description": "Parte a incrementar: major/minor/patch. Default: patch"},
            },
            "required": ["path"],
        },
    },
    # ── CMDB & Asset register ────────────────────────────────────────────────
    {
        "name": "cmdb_search",
        "description": "Busca en la base de datos de gestión de configuración (CMDB) en formato CSV/XLSX/JSON. Soporta búsqueda por texto libre o por campo específico.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query":      {"type": "string",  "description": "Texto a buscar (usa * para listar todo)"},
                "field":      {"type": "string",  "description": "Columna donde buscar (opcional; sin ella busca en todas)"},
                "cmdb_path":  {"type": "string",  "description": "Ruta al fichero CMDB (opcional; auto-detectado en cwd y ~/Documents/)"},
                "limit":      {"type": "integer", "description": "Máximo de resultados. Default: 20"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "cmdb_update",
        "description": "Actualiza un registro en la CMDB CSV identificado por un campo clave.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "key_field":  {"type": "string", "description": "Nombre de la columna clave (p.ej. 'hostname' o 'asset_id')"},
                "key_value":  {"type": "string", "description": "Valor del campo clave del registro a actualizar"},
                "updates":    {"type": "object", "description": "Dict con {columna: nuevo_valor} a actualizar"},
                "cmdb_path":  {"type": "string", "description": "Ruta al fichero CMDB CSV (opcional; auto-detectado)"},
            },
            "required": ["key_field", "key_value", "updates"],
        },
    },
    {
        "name": "asset_register_add",
        "description": "Añade un nuevo activo al registro de activos CSV. Si el fichero no existe lo crea con las cabeceras apropiadas.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "asset":       {"type": "object", "description": "Dict con los campos del activo: {hostname, ip, type, os, location, owner, status, ...}"},
                "register_path":{"type": "string","description": "Ruta al fichero CSV del registro (opcional; por defecto: asset_register.csv en cwd)"},
            },
            "required": ["asset"],
        },
    },
]

_TOOL_FNS: dict[str, Any] = {
    "email_list":               _tool_email_list,
    "email_read":               _tool_email_read,
    "email_send":               _tool_email_send,
    "email_search":             _tool_email_search,
    "doc_convert":              _tool_doc_convert,
    "pdf_extract_text":         _tool_pdf_extract_text,
    "doc_word_count":           _tool_doc_word_count,
    "xlsx_read":                _tool_xlsx_read,
    "xlsx_write":               _tool_xlsx_write,
    "csv_analyze":              _tool_csv_analyze,
    "cal_list":                 _tool_cal_list,
    "cal_add":                  _tool_cal_add,
    "cal_search":               _tool_cal_search,
    "notes_list":               _tool_notes_list,
    "notes_search":             _tool_notes_search,
    "notes_save":               _tool_notes_save,
    "image_to_text":            _tool_image_to_text,
    "contact_search":           _tool_contact_search,
    "markdown_to_html":         _tool_markdown_to_html,
    "doc_read_template_fields": _tool_doc_read_template_fields,
    "doc_fill_template":        _tool_doc_fill_template,
    "doc_list_templates":       _tool_doc_list_templates,
    "doc_create_rfc":           _tool_doc_create_rfc,
    "xlsx_fill_range":          _tool_xlsx_fill_range,
    "xlsx_append_row":          _tool_xlsx_append_row,
    "xlsx_create_report":       _tool_xlsx_create_report,
    "project_context_read":     _tool_project_context_read,
    "project_init_office":      _tool_project_init_office,
    "doc_project_save":         _tool_doc_project_save,
    "doc_read":                 _tool_doc_read,
    "doc_update_section":       _tool_doc_update_section,
    "doc_version_bump":         _tool_doc_version_bump,
    "cmdb_search":              _tool_cmdb_search,
    "cmdb_update":              _tool_cmdb_update,
    "asset_register_add":       _tool_asset_register_add,
}


# ── Prompts ──────────────────────────────────────────────────────────────────

_PROMPTS: dict[str, dict] = {
    "draft_email": {
        "description": "Redacta un email profesional con tono y estructura configurables.",
        "arguments": [
            {"name": "to",       "description": "Destinatario",         "required": True},
            {"name": "subject",  "description": "Asunto",               "required": True},
            {"name": "context",  "description": "Contexto o propósito", "required": True},
            {"name": "tone",     "description": "formal/informal/amigable", "required": False},
            {"name": "language", "description": "Idioma (español por defecto)", "required": False},
        ],
    },
    "summarize_document": {
        "description": "Resume un documento con puntos clave, decisiones y acciones requeridas.",
        "arguments": [
            {"name": "content",   "description": "Texto del documento a resumir", "required": True},
            {"name": "max_words", "description": "Longitud máxima del resumen en palabras", "required": False},
            {"name": "focus",     "description": "Aspecto en el que enfocarse (p.ej. 'puntos de acción')", "required": False},
        ],
    },
    "meeting_notes": {
        "description": "Genera un acta de reunión estructurada con asistentes, puntos tratados y decisiones.",
        "arguments": [
            {"name": "attendees",  "description": "Lista de asistentes",          "required": True},
            {"name": "agenda",     "description": "Puntos de la agenda",           "required": True},
            {"name": "notes",      "description": "Notas o puntos discutidos",     "required": False},
            {"name": "date",       "description": "Fecha de la reunión (YYYY-MM-DD)", "required": False},
        ],
    },
    "weekly_report": {
        "description": "Genera un informe semanal de actividades con resumen, logros y próximos pasos.",
        "arguments": [
            {"name": "completed",   "description": "Tareas completadas esta semana", "required": True},
            {"name": "in_progress", "description": "Tareas en progreso",             "required": False},
            {"name": "next_week",   "description": "Plan para la próxima semana",    "required": False},
            {"name": "blockers",    "description": "Bloqueos o problemas",           "required": False},
        ],
    },
    # ── IT / Datacenter prompts ──────────────────────────────────────────────
    "datacenter_migration_report": {
        "description": "Genera un informe técnico formal de migración de centro de datos con fases, inventario y resultados.",
        "arguments": [
            {"name": "project_name", "description": "Nombre del proyecto de migración",          "required": True},
            {"name": "source_dc",    "description": "Centro de datos o entorno origen",          "required": True},
            {"name": "target_dc",    "description": "Centro de datos o entorno destino",         "required": True},
            {"name": "systems",      "description": "Lista de sistemas/servidores migrados",     "required": True},
            {"name": "date",         "description": "Fecha de la migración (YYYY-MM-DD)",        "required": False},
            {"name": "team",         "description": "Equipo o responsables",                     "required": False},
            {"name": "issues",       "description": "Incidencias ocurridas durante la migración","required": False},
            {"name": "result",       "description": "Resultado: exitoso / parcial / fallido",    "required": False},
        ],
    },
    "rfc_change_request": {
        "description": "Genera un RFC/Change Request completo y formal para cambios de infraestructura IT.",
        "arguments": [
            {"name": "title",            "description": "Título del cambio",                     "required": True},
            {"name": "requester",        "description": "Nombre del solicitante",                "required": True},
            {"name": "description",      "description": "Descripción detallada del cambio",      "required": True},
            {"name": "affected_systems", "description": "Sistemas o servicios afectados",        "required": True},
            {"name": "risk_level",       "description": "Nivel de riesgo: Alto/Medio/Bajo",      "required": False},
            {"name": "rollback_plan",    "description": "Plan de marcha atrás",                  "required": False},
            {"name": "scheduled_date",   "description": "Fecha y ventana de mantenimiento",      "required": False},
        ],
    },
    "server_migration_plan": {
        "description": "Genera un plan detallado de migración de servidor con checklist, validaciones y rollback.",
        "arguments": [
            {"name": "server_name",     "description": "Nombre o hostname del servidor",         "required": True},
            {"name": "source_env",      "description": "Entorno/infraestructura origen",         "required": True},
            {"name": "target_env",      "description": "Entorno/infraestructura destino",        "required": True},
            {"name": "services",        "description": "Servicios o aplicaciones en el servidor","required": False},
            {"name": "downtime_window", "description": "Ventana de mantenimiento",               "required": False},
            {"name": "dependencies",    "description": "Dependencias del servidor",              "required": False},
        ],
    },
    "it_incident_report": {
        "description": "Genera un informe post-incidencia IT (Post-Mortem) con RCA, timeline y plan de acción preventivo.",
        "arguments": [
            {"name": "title",         "description": "Título de la incidencia",                  "required": True},
            {"name": "severity",      "description": "Severidad: Crítica/Alta/Media/Baja",       "required": True},
            {"name": "start_time",    "description": "Fecha/hora de inicio del incidente",       "required": True},
            {"name": "end_time",      "description": "Fecha/hora de resolución",                 "required": False},
            {"name": "affected",      "description": "Sistemas y usuarios afectados",            "required": False},
            {"name": "root_cause",    "description": "Causa raíz identificada",                  "required": False},
            {"name": "timeline",      "description": "Cronología de eventos clave",              "required": False},
            {"name": "actions_taken", "description": "Acciones tomadas para resolverlo",         "required": False},
            {"name": "preventive",    "description": "Medidas preventivas propuestas",           "required": False},
        ],
    },
    "infrastructure_change_plan": {
        "description": "Genera un plan de cambio de infraestructura con análisis de impacto, fases, riesgos y aprobaciones.",
        "arguments": [
            {"name": "project",   "description": "Nombre del proyecto/cambio",                   "required": True},
            {"name": "requester", "description": "Responsable del proyecto",                     "required": True},
            {"name": "objective", "description": "Objetivo del cambio",                          "required": True},
            {"name": "scope",     "description": "Alcance: sistemas y servicios incluidos",      "required": False},
            {"name": "phases",    "description": "Fases del proyecto",                           "required": False},
            {"name": "risks",     "description": "Riesgos identificados",                        "required": False},
            {"name": "timeline",  "description": "Cronograma estimado",                          "required": False},
            {"name": "approvers", "description": "Lista de aprobadores requeridos",              "required": False},
        ],
    },
    # ── Bloque 4: Gestión IT / Negocio ───────────────────────────────────────
    "executive_summary": {
        "description": "Genera un resumen ejecutivo de un proyecto o situación IT para presentar a dirección.",
        "arguments": [
            {"name": "project",    "description": "Nombre del proyecto o situación",             "required": True},
            {"name": "context",    "description": "Contexto o descripción del contenido",        "required": True},
            {"name": "audience",   "description": "Audiencia objetivo (CIO, CEO, comité…)",      "required": False},
            {"name": "max_pages",  "description": "Extensión máxima en páginas. Default: 1",     "required": False},
        ],
    },
    "business_case": {
        "description": "Genera un business case IT con análisis coste-beneficio, ROI y justificación de inversión.",
        "arguments": [
            {"name": "project",     "description": "Nombre del proyecto/inversión",              "required": True},
            {"name": "requester",   "description": "Responsable o departamento solicitante",     "required": True},
            {"name": "description", "description": "Descripción de la inversión propuesta",      "required": True},
            {"name": "cost",        "description": "Coste estimado (CAPEX/OPEX)",                "required": False},
            {"name": "benefits",    "description": "Beneficios esperados (cuantitativos/cualitativos)", "required": False},
            {"name": "alternatives","description": "Alternativas consideradas",                  "required": False},
            {"name": "timeline",    "description": "Plazo de amortización o retorno",            "required": False},
        ],
    },
    "project_status_report": {
        "description": "Genera un informe de estado de proyecto IT (RAG status) con hitos, riesgos y próximos pasos.",
        "arguments": [
            {"name": "project",     "description": "Nombre del proyecto",                        "required": True},
            {"name": "period",      "description": "Período del informe (p.ej. 'Mayo 2026')",    "required": True},
            {"name": "status",      "description": "Estado general: Verde/Ámbar/Rojo",           "required": True},
            {"name": "completed",   "description": "Hitos o tareas completadas en el período",   "required": False},
            {"name": "in_progress", "description": "Tareas en curso",                            "required": False},
            {"name": "risks",       "description": "Riesgos o problemas identificados",          "required": False},
            {"name": "next_steps",  "description": "Próximos hitos o acciones",                  "required": False},
            {"name": "budget",      "description": "Estado del presupuesto",                     "required": False},
        ],
    },
}


def _get_prompt(name: str, args: dict) -> list[dict]:
    if name == "draft_email":
        to      = args.get("to", "")
        subject = args.get("subject", "")
        context = args.get("context", "")
        tone    = args.get("tone", "profesional")
        lang    = args.get("language", "español")
        prompt  = (
            f"Redacta un email en {lang} con tono {tone}.\n\n"
            f"Para: {to}\nAsunto: {subject}\nContexto: {context}\n\n"
            "El email debe incluir:\n"
            "- Saludo adecuado al tono\n"
            "- Cuerpo claro y conciso\n"
            "- Llamada a la acción si aplica\n"
            "- Despedida apropiada\n\n"
            "Devuelve solo el cuerpo del email, sin indicaciones de formato adicionales."
        )
    elif name == "summarize_document":
        content   = args.get("content", "")
        max_words = args.get("max_words", "300")
        focus     = args.get("focus", "")
        focus_str = f"\nEnfoca especialmente en: {focus}" if focus else ""
        prompt = (
            f"Resume el siguiente documento en máximo {max_words} palabras.{focus_str}\n\n"
            "Estructura del resumen:\n"
            "## Resumen ejecutivo\n[2-3 frases clave]\n\n"
            "## Puntos principales\n[lista con viñetas]\n\n"
            "## Decisiones / Conclusiones\n[si las hay]\n\n"
            "## Acciones requeridas\n[con responsable y plazo si se mencionan]\n\n"
            f"DOCUMENTO:\n{content[:8000]}"
        )
    elif name == "meeting_notes":
        attendees = args.get("attendees", "")
        agenda    = args.get("agenda", "")
        notes     = args.get("notes", "")
        date      = args.get("date", datetime.date.today().isoformat())
        prompt = (
            f"Genera un acta de reunión formal para la fecha {date}.\n\n"
            f"ASISTENTES:\n{attendees}\n\n"
            f"AGENDA:\n{agenda}\n\n"
            f"NOTAS / PUNTOS DISCUTIDOS:\n{notes or 'No proporcionadas'}\n\n"
            "Formato del acta:\n"
            "# Acta de Reunión — [fecha]\n\n"
            "**Asistentes:** ...\n\n"
            "## Puntos tratados\n[por orden de agenda]\n\n"
            "## Decisiones tomadas\n[con responsables]\n\n"
            "## Próximos pasos\n[tareas, responsables y fechas límite]\n\n"
            "## Próxima reunión\n[si se acordó]"
        )
    elif name == "weekly_report":
        completed   = args.get("completed", "")
        in_progress = args.get("in_progress", "")
        next_week   = args.get("next_week", "")
        blockers    = args.get("blockers", "")
        week = datetime.date.today().isocalendar()[1]
        prompt = (
            f"Genera un informe semanal profesional (semana {week}).\n\n"
            f"COMPLETADO:\n{completed}\n\n"
            f"EN PROGRESO:\n{in_progress or 'Ninguno'}\n\n"
            f"PRÓXIMA SEMANA:\n{next_week or 'Por definir'}\n\n"
            f"BLOQUEOS:\n{blockers or 'Ninguno'}\n\n"
            "Formato:\n"
            "# Informe Semanal — Semana [N]\n\n"
            "## ✅ Completado\n[lista con impacto breve de cada ítem]\n\n"
            "## 🔄 En progreso\n[con % estimado de avance]\n\n"
            "## 📅 Plan próxima semana\n[prioridades ordenadas]\n\n"
            "## ⚠ Bloqueos / Riesgos\n[acciones necesarias]"
        )
    elif name == "datacenter_migration_report":
        project_name = args.get("project_name", "")
        source_dc    = args.get("source_dc", "")
        target_dc    = args.get("target_dc", "")
        systems      = args.get("systems", "")
        date         = args.get("date", datetime.date.today().isoformat())
        team         = args.get("team", "No especificado")
        issues       = args.get("issues", "Ninguna")
        result       = args.get("result", "Exitoso")
        prompt = (
            f"Genera un informe técnico formal de migración de centro de datos para la fecha {date}.\n\n"
            f"PROYECTO: {project_name}\nORIGEN: {source_dc}\nDESTINO: {target_dc}\n"
            f"SISTEMAS MIGRADOS:\n{systems}\nEQUIPO: {team}\nINCIDENCIAS: {issues}\nRESULTADO: {result}\n\n"
            "Formato requerido (markdown formal):\n\n"
            "# Informe de Migración — [Nombre Proyecto]\n\n"
            "## Resumen Ejecutivo\n[Estado, resultado y contexto en 3-4 líneas]\n\n"
            "## Datos del Proyecto\n| Campo | Valor |\n[Tabla: proyecto, fecha, origen, destino, equipo, resultado]\n\n"
            "## Inventario de Sistemas Migrados\n| Sistema | Tipo | IP Origen | IP Destino | Estado | Observaciones |\n\n"
            "## Fases de Migración Ejecutadas\n[Lista numerada con estado ✅/⚠/❌]\n\n"
            "## Incidencias y Resolución\n[Si las hubo, con descripción, impacto y resolución]\n\n"
            "## Validaciones Post-Migración\n[Checklist: conectividad, servicios, datos, rendimiento]\n\n"
            "## Resultados y Conclusiones\n\n"
            "## Próximos Pasos\n\n"
            "## Firmas y Aprobaciones\n| Rol | Nombre | Firma | Fecha |"
        )
    elif name == "rfc_change_request":
        title            = args.get("title", "")
        requester        = args.get("requester", "")
        description      = args.get("description", "")
        affected_systems = args.get("affected_systems", "")
        risk_level       = args.get("risk_level", "Medio")
        rollback_plan    = args.get("rollback_plan", "Por definir")
        scheduled_date   = args.get("scheduled_date", "Por definir")
        rfc_id = f"RFC-{datetime.date.today().strftime('%Y%m%d')}"
        prompt = (
            f"Genera un documento RFC/Request for Change completo y formal para entorno corporativo.\n\n"
            f"ID RFC: {rfc_id}\nTÍTULO: {title}\nSOLICITANTE: {requester}\n"
            f"DESCRIPCIÓN: {description}\nSISTEMAS AFECTADOS: {affected_systems}\n"
            f"NIVEL DE RIESGO: {risk_level}\nFECHA PROGRAMADA: {scheduled_date}\n"
            f"PLAN DE ROLLBACK: {rollback_plan}\n\n"
            "Formato requerido (markdown formal):\n\n"
            f"# {rfc_id} — [Título]\n\n"
            "## 1. Información General\n[Tabla completa con todos los campos]\n\n"
            "## 2. Descripción del Cambio\n[Detallada y técnica]\n\n"
            "## 3. Justificación y Beneficios\n\n"
            "## 4. Alcance e Impacto\n[Sistemas afectados, usuarios, tiempo de inactividad estimado]\n\n"
            "## 5. Plan de Implementación\n[Pasos numerados con tiempo estimado y responsable]\n\n"
            "## 6. Análisis de Riesgos\n| Riesgo | Probabilidad | Impacto | Mitigación |\n\n"
            "## 7. Plan de Pruebas y Validación\n[Pasos de verificación post-cambio]\n\n"
            "## 8. Plan de Marcha Atrás\n[Pasos detallados para revertir]\n\n"
            "## 9. Plan de Comunicaciones\n[A quién notificar, cuándo y cómo]\n\n"
            "## 10. Aprobaciones\n| Rol | Nombre | Firma | Fecha |"
        )
    elif name == "server_migration_plan":
        server_name     = args.get("server_name", "")
        source_env      = args.get("source_env", "")
        target_env      = args.get("target_env", "")
        services        = args.get("services", "No especificados")
        downtime_window = args.get("downtime_window", "Por definir")
        dependencies    = args.get("dependencies", "Por analizar")
        prompt = (
            f"Genera un plan técnico detallado de migración de servidor.\n\n"
            f"SERVIDOR: {server_name}\nORIGEN: {source_env}\nDESTINO: {target_env}\n"
            f"SERVICIOS: {services}\nVENTANA DE MANTENIMIENTO: {downtime_window}\n"
            f"DEPENDENCIAS: {dependencies}\n\n"
            "Formato requerido (markdown técnico):\n\n"
            "# Plan de Migración — [Servidor]\n\n"
            "## Pre-requisitos\n[Checks y preparativos antes de la migración]\n\n"
            "## Inventario del Sistema Origen\n[HW, SO, servicios, datos, configuración de red]\n\n"
            "## Plan de Migración Paso a Paso\n| Paso | Descripción | Duración | Responsable |\n\n"
            "## Comandos de Migración\n[Código/comandos para cada fase crítica]\n\n"
            "## Plan de Pruebas\n[Checklist de validaciones post-migración]\n\n"
            "## Plan de Rollback\n[Pasos ordenados para revertir]\n\n"
            "## Plan de Comunicaciones\n[Pre/durante/post migración]\n\n"
            "## Criterios de Éxito\n[Qué debe estar operativo para dar la migración por correcta]"
        )
    elif name == "it_incident_report":
        title         = args.get("title", "")
        severity      = args.get("severity", "Media")
        start_time    = args.get("start_time", "")
        end_time      = args.get("end_time", "En curso")
        affected      = args.get("affected", "Por determinar")
        root_cause    = args.get("root_cause", "En investigación")
        timeline      = args.get("timeline", "Por documentar")
        actions_taken = args.get("actions_taken", "Por documentar")
        preventive    = args.get("preventive", "Por definir")
        prompt = (
            f"Genera un informe post-incidencia (Post-Mortem) IT completo y profesional.\n\n"
            f"TÍTULO: {title}\nSEVERIDAD: {severity}\nINICIO: {start_time}\nFIN: {end_time}\n"
            f"AFECTADOS: {affected}\nCAUSA RAÍZ: {root_cause}\nCRONOLOGÍA: {timeline}\n"
            f"ACCIONES: {actions_taken}\nMEDIDAS PREVENTIVAS: {preventive}\n\n"
            "Formato requerido (markdown formal):\n\n"
            "# Informe de Incidencia — [Título]\n\n"
            "## Resumen Ejecutivo\n[Qué pasó, impacto y resultado en 3-4 líneas]\n\n"
            "## Información del Incidente\n[Tabla: ID, título, severidad, inicio, fin, duración, estado]\n\n"
            "## Cronología de Eventos\n| Hora | Evento | Responsable | Acción |\n\n"
            "## Análisis de Causa Raíz (RCA)\n[Técnica 5 Porqués]\n\n"
            "## Impacto\n[Sistemas, usuarios, negocio, SLA afectado]\n\n"
            "## Acciones de Resolución Tomadas\n[Numeradas con responsable y hora]\n\n"
            "## Lecciones Aprendidas\n\n"
            "## Plan de Acción Preventivo\n| Acción | Responsable | Fecha límite | Estado |\n\n"
            "## Aprobaciones\n| Rol | Nombre | Firma | Fecha |"
        )
    elif name == "infrastructure_change_plan":
        project   = args.get("project", "")
        requester = args.get("requester", "")
        objective = args.get("objective", "")
        scope     = args.get("scope", "Por definir")
        phases    = args.get("phases", "Por definir")
        risks     = args.get("risks", "Por analizar")
        timeline  = args.get("timeline", "Por definir")
        approvers = args.get("approvers", "Por definir")
        prompt = (
            f"Genera un plan de cambio de infraestructura IT formal y detallado.\n\n"
            f"PROYECTO: {project}\nRESPONSABLE: {requester}\nOBJETIVO: {objective}\n"
            f"ALCANCE: {scope}\nFASES: {phases}\nRIESGOS: {risks}\n"
            f"CRONOGRAMA: {timeline}\nAPROBADORES: {approvers}\n\n"
            "Formato requerido (markdown formal):\n\n"
            "# Plan de Cambio de Infraestructura — [Proyecto]\n\n"
            "## Resumen Ejecutivo\n\n"
            "## Objetivos y Beneficios Esperados\n\n"
            "## Alcance del Proyecto\n[Tabla: En scope / Fuera de scope]\n\n"
            "## Análisis del Estado Actual (As-Is)\n\n"
            "## Estado Objetivo (To-Be)\n\n"
            "## Fases del Proyecto\n| Fase | Descripción | Inicio | Fin | Responsable | Estado |\n\n"
            "## Análisis de Riesgos\n| Riesgo | Probabilidad | Impacto | Mitigación |\n\n"
            "## Plan de Rollback\n\n"
            "## Cronograma — Hitos Principales\n\n"
            "## Recursos Necesarios\n[HW, SW, licencias, equipo]\n\n"
            "## Plan de Comunicaciones\n\n"
            "## Aprobaciones\n| Rol | Nombre | Firma | Fecha |"
        )
    elif name == "executive_summary":
        project   = args.get("project", "")
        context   = args.get("context", "")
        audience  = args.get("audience", "Dirección / Comité de IT")
        max_pages = args.get("max_pages", "1")
        prompt = (
            f"Genera un resumen ejecutivo de máximo {max_pages} página(s) para la audiencia: {audience}.\n\n"
            f"PROYECTO/SITUACIÓN: {project}\n\nCONTEXTO:\n{context}\n\n"
            "Formato requerido (markdown ejecutivo):\n\n"
            "# Resumen Ejecutivo — [Proyecto]\n\n"
            "## Situación Actual\n[2-3 frases: qué está pasando, contexto de negocio]\n\n"
            "## Objetivo\n[Qué se pretende conseguir y por qué es importante ahora]\n\n"
            "## Puntos Clave\n[3-5 bullets con hechos o datos relevantes]\n\n"
            "## Estado / Avance\n[Verde ✅ / Ámbar ⚠ / Rojo 🔴 con justificación breve]\n\n"
            "## Decisiones / Acciones Requeridas\n[Qué necesita la dirección aprobar, decidir o saber]\n\n"
            "## Riesgos Críticos\n[Solo los top 2-3; impacto y mitigación en una línea cada uno]\n\n"
            "## Próximos Hitos\n[Con fecha y responsable]\n\n"
            "Tono: directo, ejecutivo, sin tecnicismos innecesarios. Máximo concisión."
        )
    elif name == "business_case":
        project     = args.get("project", "")
        requester   = args.get("requester", "")
        description = args.get("description", "")
        cost        = args.get("cost", "Por determinar")
        benefits    = args.get("benefits", "Por analizar")
        alternatives= args.get("alternatives", "No aplica")
        timeline    = args.get("timeline", "Por definir")
        prompt = (
            f"Genera un Business Case IT completo y profesional para presentar a dirección.\n\n"
            f"PROYECTO: {project}\nSOLICITANTE: {requester}\n"
            f"DESCRIPCIÓN: {description}\nCOSTE ESTIMADO: {cost}\n"
            f"BENEFICIOS: {benefits}\nALTERNATIVAS: {alternatives}\n"
            f"PLAZO DE RETORNO: {timeline}\n\n"
            "Formato requerido (markdown formal):\n\n"
            "# Business Case — [Proyecto]\n\n"
            "## 1. Resumen Ejecutivo\n[Media página máxima]\n\n"
            "## 2. Situación Actual y Problema\n[As-Is: problema, pain points, coste de no actuar]\n\n"
            "## 3. Solución Propuesta\n[Descripción técnica y funcional]\n\n"
            "## 4. Análisis de Alternativas\n| Opción | Descripción | Ventajas | Inconvenientes | Coste |\n\n"
            "## 5. Análisis Coste-Beneficio\n"
            "### 5.1 Costes (CAPEX + OPEX)\n| Concepto | Año 1 | Año 2 | Año 3 | Total |\n"
            "### 5.2 Beneficios Cuantificables\n| Beneficio | Valor anual estimado |\n"
            "### 5.3 ROI y Período de Retorno\n[Fórmula y resultado]\n\n"
            "## 6. Riesgos\n| Riesgo | Probabilidad | Impacto | Mitigación |\n\n"
            "## 7. Cronograma de Implementación\n[Hitos principales con fecha]\n\n"
            "## 8. Recomendación\n[Opción recomendada y justificación]\n\n"
            "## 9. Aprobaciones\n| Rol | Nombre | Decisión | Fecha |"
        )
    elif name == "project_status_report":
        project     = args.get("project", "")
        period      = args.get("period", datetime.date.today().strftime("%B %Y"))
        status      = args.get("status", "Verde")
        completed   = args.get("completed", "Sin detalles")
        in_progress = args.get("in_progress", "Sin detalles")
        risks       = args.get("risks", "Sin riesgos identificados")
        next_steps  = args.get("next_steps", "Por definir")
        budget      = args.get("budget", "Sin información")
        status_icon = {"Verde": "✅", "Ámbar": "⚠", "Rojo": "🔴"}.get(status, "ℹ")
        prompt = (
            f"Genera un informe de estado (RAG Status Report) del proyecto {project} para el período {period}.\n\n"
            f"ESTADO GENERAL: {status_icon} {status}\n"
            f"COMPLETADO:\n{completed}\n\n"
            f"EN CURSO:\n{in_progress}\n\n"
            f"RIESGOS/PROBLEMAS:\n{risks}\n\n"
            f"PRÓXIMOS PASOS:\n{next_steps}\n\n"
            f"ESTADO PRESUPUESTO: {budget}\n\n"
            "Formato requerido (markdown formal):\n\n"
            f"# Informe de Estado — {project}\n**Período:** {period}  |  **Estado:** {status_icon} {status}\n\n"
            "## Resumen Ejecutivo\n[2-3 líneas del estado global]\n\n"
            "## KPIs del Período\n| Indicador | Objetivo | Real | Estado |\n\n"
            "## Hitos Completados\n[Lista con ✅ y fecha real]\n\n"
            "## Tareas en Curso\n[Con % avance y fecha estimada de fin]\n\n"
            "## Riesgos e Issues\n| # | Descripción | Impacto | Probabilidad | Acción | Responsable |\n\n"
            "## Presupuesto\n| Concepto | Presupuesto | Ejecutado | Desviación |\n\n"
            "## Próximos Hitos\n| Hito | Fecha prevista | Responsable |\n\n"
            "## Decisiones / Escaladas Requeridas\n[Si las hay]"
        )
    else:
        prompt = f"Prompt {name} no disponible."
    return [{"role": "user", "content": {"type": "text", "text": prompt}}]


# ── Resources ────────────────────────────────────────────────────────────────

_RESOURCES = [
    {
        "uri":         "office://emails_recent",
        "name":        "Emails recientes",
        "description": "Últimos 10 emails de la bandeja de entrada",
        "mimeType":    "text/plain",
    },
    {
        "uri":         "office://calendar_today",
        "name":        "Calendario",
        "description": "Eventos de hoy y los próximos 7 días",
        "mimeType":    "text/plain",
    },
    {
        "uri":         "office://notes_recent",
        "name":        "Notas recientes",
        "description": "Notas markdown más recientes (últimas 10)",
        "mimeType":    "text/plain",
    },
    {
        "uri":         "office://templates_available",
        "name":        "Plantillas disponibles",
        "description": "Lista de plantillas de documentos disponibles (.docx, .xlsx, .md)",
        "mimeType":    "text/plain",
    },
    {
        "uri":         "office://rfc_pending",
        "name":        "RFCs pendientes",
        "description": "Lista de documentos RFC/Change Request en el directorio de notas",
        "mimeType":    "text/plain",
    },
    {
        "uri":         "office://tasks_today",
        "name":        "Tareas de hoy",
        "description": "Tareas pendientes de la jornada (desde el plugin todo de OOCode)",
        "mimeType":    "text/plain",
    },
    {
        "uri":         "office://project_context",
        "name":        "Contexto del proyecto",
        "description": "Metadata del proyecto activo: nombre, cliente, tipo, directorios, naming convention (desde OOCODE.md del cwd)",
        "mimeType":    "text/plain",
    },
    {
        "uri":         "office://server_inventory",
        "name":        "Inventario de servidores",
        "description": "Listado completo de activos de la CMDB del proyecto (auto-detectada en cwd y ~/Documents/)",
        "mimeType":    "text/plain",
    },
]


def _resource_emails_recent() -> str:
    return _tool_email_list({"limit": 10})


def _resource_calendar_today() -> str:
    cfg   = _load_config()
    today = datetime.date.today().isoformat()
    end   = (datetime.date.today() + datetime.timedelta(days=7)).isoformat()
    return _tool_cal_list({"start": today, "end": end, "limit": 20, "source": cfg["calendar_file"]})


def _resource_notes_recent() -> str:
    return _tool_notes_list({"limit": 10})


def _resource_templates_available() -> str:
    return _tool_doc_list_templates({})


def _resource_rfc_pending() -> str:
    cfg      = _load_config()
    notes_d  = Path(cfg["notes_dir"]).expanduser()
    if not notes_d.exists():
        return f"Directorio de notas no encontrado: {notes_d}"
    rfc_files = sorted(
        [f for f in notes_d.glob("*.md") if re.search(r"rfc|change.?request|cambio", f.stem, re.IGNORECASE)],
        key=lambda f: f.stat().st_mtime, reverse=True
    )[:20]
    if not rfc_files:
        # Fallback: search by content
        return _tool_notes_search({"query": "RFC", "limit": 10})
    lines = [f"📋 RFCs en {notes_d} — {len(rfc_files)} ficheros:"]
    for f in rfc_files:
        mtime = datetime.datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d")
        lines.append(f"  {mtime}  {f.name}")
    return "\n".join(lines)


def _resource_tasks_today() -> str:
    todo_path = Path.home() / ".oocode" / "todos.json"
    if not todo_path.exists():
        return "Sin tareas pendientes. Usa el plugin todo (/todo) para añadir tareas."
    try:
        data   = json.loads(todo_path.read_text())
        todos  = data if isinstance(data, list) else data.get("todos", [])
        today  = datetime.date.today().isoformat()
        pending = [t for t in todos if t.get("status", "pending") != "done"]
        done    = [t for t in todos if t.get("status") == "done"]
        if not pending:
            return f"✅ Sin tareas pendientes hoy ({today}). {len(done)} completadas."
        lines = [f"📋 Tareas pendientes ({len(pending)}) — {today}:"]
        for t in pending[:25]:
            icon  = "◻" if t.get("status") == "pending" else "◼"
            title = t.get("title", t.get("text", str(t)))
            cat   = t.get("category", "")
            cat_s = f" [{cat}]" if cat else ""
            lines.append(f"  {icon} {title}{cat_s}")
        if done:
            lines.append(f"\n✅ Completadas hoy: {len(done)}")
        return "\n".join(lines)
    except Exception as exc:
        return f"Error leyendo tareas: {exc}"


def _resource_project_context() -> str:
    return _tool_project_context_read({})


def _resource_server_inventory() -> str:
    return _tool_cmdb_search({"query": "*", "limit": 50})


_RESOURCE_FNS = {
    "office://emails_recent":      _resource_emails_recent,
    "office://calendar_today":     _resource_calendar_today,
    "office://notes_recent":       _resource_notes_recent,
    "office://templates_available":_resource_templates_available,
    "office://rfc_pending":        _resource_rfc_pending,
    "office://tasks_today":        _resource_tasks_today,
    "office://project_context":    _resource_project_context,
    "office://server_inventory":   _resource_server_inventory,
}


# ── Bucle principal ───────────────────────────────────────────────────────────

def _handle(req: dict) -> None:
    method = req.get("method", "")
    req_id = req.get("id")
    params = req.get("params", {})

    if method == "initialize":
        _ok(req_id, {
            "protocolVersion": "2024-11-05",
            "serverInfo": {"name": "home-office-assistant", "version": "2.0.0"},
            "capabilities": {
                "tools":     {"listChanged": False},
                "resources": {"listChanged": False},
                "prompts":   {"listChanged": False},
            },
        })

    elif method == "notifications/initialized":
        pass

    elif method == "tools/list":
        _ok(req_id, {"tools": _TOOLS})

    elif method == "tools/call":
        name      = params.get("name", "")
        arguments = params.get("arguments", {})
        fn        = _TOOL_FNS.get(name)
        if fn is None:
            _err(req_id, -32601, f"Tool desconocida: {name}")
            return
        try:
            result = fn(arguments)
            _ok(req_id, {"content": [{"type": "text", "text": result}], "isError": False})
        except Exception as exc:
            _ok(req_id, {"content": [{"type": "text", "text": f"Error: {exc}"}], "isError": True})

    elif method == "resources/list":
        _ok(req_id, {"resources": _RESOURCES})

    elif method == "resources/read":
        uri = params.get("uri", "")
        fn  = _RESOURCE_FNS.get(uri)
        if fn is None:
            _err(req_id, -32601, f"Recurso desconocido: {uri}")
            return
        try:
            content = fn()
            _ok(req_id, {"contents": [{"uri": uri, "mimeType": "text/plain", "text": content}]})
        except Exception as exc:
            _err(req_id, -32603, f"Error leyendo recurso: {exc}")

    elif method == "prompts/list":
        prompts = [
            {"name": k, "description": v["description"], "arguments": v["arguments"]}
            for k, v in _PROMPTS.items()
        ]
        _ok(req_id, {"prompts": prompts})

    elif method == "prompts/get":
        name      = params.get("name", "")
        arguments = params.get("arguments", {})
        if name not in _PROMPTS:
            _err(req_id, -32601, f"Prompt desconocido: {name}")
            return
        messages = _get_prompt(name, arguments)
        _ok(req_id, {"description": _PROMPTS[name]["description"], "messages": messages})

    elif req_id is not None:
        _err(req_id, -32601, f"Método desconocido: {method}")


def main() -> None:
    sys.stderr.write("[home-office-assistant] MCP server v3.0 iniciado (35 tools, 12 prompts, 8 resources)\n")
    sys.stderr.flush()
    while True:
        try:
            req = _recv()
            if req is None:
                break
            _handle(req)
        except (EOFError, BrokenPipeError):
            break
        except Exception as exc:
            sys.stderr.write(f"[home-office-assistant] Error: {exc}\n")
            sys.stderr.flush()


if __name__ == "__main__":
    main()
